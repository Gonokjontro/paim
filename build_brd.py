from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/ashekurrahman/Documents/PAIM/deliverables/Personal_AI_Subscription_Management_BRD_v1.0.docx")

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
WHITE = "FFFFFF"
GREEN = "216E39"
PALE_GREEN = "EAF5EE"
AMBER = "7A5A00"
PALE_AMBER = "FFF6DD"
RED = "9B1C1C"
PALE_RED = "FCE8E6"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, italic=None, color=DARK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction, placeholder="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def paragraph_border_bottom(paragraph, color=BLUE, size="12", space="6"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_real_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs or [0]) + 1
    bullet_num_id = max(existing_num or [0]) + 1
    decimal_num_id = bullet_num_id + 1

    def abstract(num_fmt, text, font=None):
        abs_num = OxmlElement("w:abstractNum")
        abs_num.set(qn("w:abstractNumId"), str(abstract_id if num_fmt == "bullet" else abstract_id + 1))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abs_num.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, lvl_text, jc, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abs_num.append(lvl)
        numbering.append(abs_num)

    abstract("bullet", "•", "Symbol")
    abstract("decimal", "%1.")

    for num_id, abs_id in ((bullet_num_id, abstract_id), (decimal_num_id, abstract_id + 1)):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)
    return bullet_num_id, decimal_num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def clone_num_instance(doc, base_num_id):
    numbering = doc.part.numbering_part.element
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_num_id = max(existing_num or [0]) + 1
    base = None
    for node in numbering.findall(qn("w:num")):
        if int(node.get(qn("w:numId"))) == base_num_id:
            base = node
            break
    if base is None:
        return base_num_id
    abstract_ref = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), abstract_ref)
    num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return new_num_id


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

bullet_num_id, decimal_num_id = add_real_numbering(doc)


def add_p(text="", bold_lead=None, italic=False, color=DARK, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=10.5, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=10.5, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, italic=italic, color=color)
    return p


def add_bullets(items):
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        r = p.add_run("•  " + text)
        set_run_font(r, size=10.5)


def add_numbers(items):
    list_num_id = clone_num_instance(doc, decimal_num_id)
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.10
        apply_num(p, list_num_id)
        r = p.add_run(text)
        set_run_font(r, size=10.5)


def add_compact_contents(items):
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.30)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        r = p.add_run("•  " + text)
        set_run_font(r, size=8.8)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_table(headers, rows, widths, font_size=8.5, header_fill=NAVY,
              alignments=None, row_fills=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_row_cant_split(hdr)
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        set_table_geometry(table, widths)
        if row_fills and ridx < len(row_fills) and row_fills[ridx]:
            for cell in cells:
                set_cell_shading(cell, row_fills[ridx])
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if alignments:
                p.alignment = alignments[idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(label, text, fill=PALE_BLUE, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement("w:" + edge)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), fill)
        p_bdr.append(border)
    p_pr.append(p_bdr)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=9.5, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5, color=DARK)
    return p


def page_break():
    doc.add_page_break()


def section_break():
    doc.add_section(WD_SECTION.NEW_PAGE)


def keep_table_rows(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True


def add_requirement_group(title, rows):
    add_heading(title, 2)
    table = add_table(
        ["ID", "Requirement", "Priority", "Acceptance summary"],
        rows,
        [850, 4450, 900, 3160],
        font_size=8.2,
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
    )
    keep_table_rows(table)
    return table


def add_use_case(uc_id, name, actor, goal, preconditions, trigger, main_flow,
                 alternate, postconditions, requirements):
    add_heading(f"{uc_id} — {name}", 2)
    meta_table = add_table(
        ["Field", "Detail"],
        [
            ("Primary actor", actor),
            ("Goal", goal),
            ("Preconditions", preconditions),
            ("Trigger", trigger),
            ("Postconditions", postconditions),
            ("Related requirements", requirements),
        ],
        [1800, 7560],
        font_size=8.6,
        header_fill=DARK_BLUE,
    )
    for row in meta_table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    add_heading("Main success flow", 3)
    add_numbers(main_flow)
    if alternate:
        add_heading("Alternate / exception flows", 3)
        add_bullets(alternate)


def add_test_table(rows):
    table = add_table(
        ["Test ID", "Use case / area", "Scenario and steps", "Expected result", "Type"],
        rows,
        [800, 1450, 3200, 3110, 800],
        font_size=7.7,
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    keep_table_rows(table)
    return table


# Running header and footer.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run("PERSONAL AI SUBSCRIPTION MANAGEMENT")
set_run_font(hr, size=8, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("Business Requirements Document  |  Page ")
set_run_font(fr, size=8, color=MID_GRAY)
add_field(fp, "PAGE", "1")


# Cover
for _ in range(3):
    add_p("", after=8)
k = add_p("BUSINESS REQUIREMENTS DOCUMENT", color=BLUE, after=8)
set_run_font(k.runs[0], size=11, bold=True, color=BLUE)
title = add_p("Personal AI Subscription Management Platform", after=8)
set_run_font(title.runs[0], size=27, bold=True, color=NAVY)
subtitle = add_p(
    "Centralized tool inventory, recurring and on-demand cost control, token usage, "
    "payment-source governance, budgets, renewals, and alerts",
    color=MID_GRAY,
    after=18,
)
set_run_font(subtitle.runs[0], size=13, color=MID_GRAY)
rule = add_p("", after=16)
paragraph_border_bottom(rule, color=BLUE, size="18", space="1")

add_table(
    ["Document", "Value"],
    [
        ("Version", "1.0"),
        ("Status", "Baseline for stakeholder review and solution delivery"),
        ("Prepared for", "Personal AI Subscription Management initiative"),
        ("Target technology", "PHP Laravel, MySQL, Metronic UI template"),
        ("Prepared on", "29 July 2026"),
        ("Classification", "Internal / Confidential"),
    ],
    [2300, 7060],
    font_size=9.3,
    header_fill=NAVY,
)
add_callout(
    "PURPOSE",
    "Define the complete business, functional, UX, data, security, reporting, and acceptance requirements "
    "for a compact, mobile-responsive system that gives users reliable control of all AI and related software subscriptions.",
)

page_break()

# 16 delivery and rollout
add_heading("16. Delivery Phasing, Migration and Rollout", 1)
add_heading("16.1 Recommended Release Plan", 2)
add_table(
    ["Phase", "Scope", "Exit criteria"],
    [
        ("Phase 0 — Foundation", "Architecture, identity, workspace boundary, design system, reference data, CI/CD, logging, security baseline.", "Technical runway, security model and UI component baseline approved."),
        ("Phase 1 — Core MVP", "Tool/vendor/type, recurring subscriptions, masked payment accounts, ledger, dashboard, monthly/annual normalization, targets, in-app/email alerts, audit.", "Must-have recurring-cost journey and core acceptance tests pass."),
        ("Phase 2 — Variable Cost", "Token packages, meters/rates, usage import, on-demand/hybrid forecasting, balances, stale-data and overage alerts.", "Reference pricing/usage calculations and token tests pass."),
        ("Phase 3 — Control & Insight", "Reconciliation, period close, plan change/cancellation evidence, advanced reports, saved views, import/export, payment reassignment.", "End-to-end monthly close and renewal-control UAT passes."),
        ("Phase 4 — Integrations & Optimization", "Approved vendor/FX/connectors, scheduled reports, recommendations, optional team/workspace expansion.", "Connector resilience, privacy authorization and performance gates pass."),
    ],
    [1700, 4800, 2860],
    font_size=8.1,
)
add_heading("16.2 MVP Boundary", 2)
add_p(
    "MVP includes all Must requirements except an explicitly approved connector-dependent item that has a complete manual/import fallback. "
    "It must support personal production use with recurring, prepaid, on-demand and hybrid subscriptions; payment aliases; accurate cost ledger; "
    "targets; alerts; dashboard; responsive UI; exports; audit; security; and operational recovery."
)
add_heading("16.3 Migration and Initial Data Load", 2)
add_numbers([
    "Inventory source files, columns, currencies, date conventions, duplicates, sensitive data and ownership.",
    "Map vendors, tools, categories, subscription types, statuses, payment aliases, cost history and usage.",
    "Clean prohibited secrets/full payment numbers before files enter the application migration process.",
    "Run dry-run import with row-level validation, duplicate rules, currency conversion and sample reconciliation.",
    "Obtain business approval of counts, active status, next renewals, normalized costs, payment links and opening targets.",
    "Execute production import in controlled window, preserve batch evidence, and reconcile dashboard/report totals.",
    "Retain source files only under the approved secure retention policy; remove staging copies after acceptance.",
])
add_heading("16.4 Rollout and Change Management", 2)
add_bullets([
    "Pilot with representative monthly, annual, token, on-demand, hybrid, multi-currency, and expiring-payment scenarios.",
    "Provide a quick-start guide, calculation glossary, import template, notification guide, and renewal/cancellation checklist.",
    "Use a brief parallel reconciliation period against source statements/spreadsheets before declaring the platform authoritative.",
    "Measure onboarding completion, unmatched entries, stale usage, alert action rate, and user-reported calculation questions.",
    "Hold a post-launch review after the first complete reporting cycle and prioritize corrections before optional integrations.",
])

# 17 risks, dependencies, decisions
add_heading("17. Risks, Dependencies, Decisions and Sign-off", 1)
add_heading("17.1 Risk Register", 2)
add_table(
    ["ID", "Risk", "Impact", "Mitigation / control", "Owner"],
    [
        ("R-01", "Incorrect cost, FX, tier or forecast calculation", "High", "Central tested domain services, reference examples, explainable totals, reconciliation and UAT.", "Tech Lead / QA"),
        ("R-02", "Users enter incomplete or stale usage/cost data", "High", "Freshness indicators, reminders, imports/connectors, exception queues and disclosed forecast basis.", "Product Owner"),
        ("R-03", "Sensitive payment or credential data is collected", "Critical", "Data minimization, prohibited fields, input warnings, masking, secure secrets for approved connectors only.", "Security"),
        ("R-04", "Alert fatigue causes important events to be ignored", "High", "Severity, dedupe, digest, quiet hours, escalation, acknowledgement analytics and sane defaults.", "Product Owner"),
        ("R-05", "Vendor API changes or outages", "Medium", "Adapter isolation, monitoring, retries, circuit breaker, stale status and manual/import fallback.", "Tech Lead"),
        ("R-06", "Metronic customization creates upgrade/accessibility debt", "Medium", "Application wrapper layer, token mapping, remove demos, accessibility tests and license/dependency inventory.", "UX / Tech Lead"),
        ("R-07", "Large ledger/usage volume slows dashboards", "Medium", "Indexes, query tests, async summaries, caching with tenant keys, performance gates and archive policy.", "DB / Tech Lead"),
        ("R-08", "Accidental deletion or historical overwrite", "High", "Soft delete, reversal workflow, plan versions, confirmation/impact preview, audit and backup restore drills.", "Tech Lead"),
        ("R-09", "Ambiguous ownership of targets and renewals", "Medium", "Mandatory owner/assignee, role matrix, calendar, escalation and overdue workflow.", "Product Owner"),
        ("R-10", "Multi-currency totals are misunderstood", "Medium", "Original/base display, stored rate/source/date, freshness warnings and calculation explanation.", "Product Owner"),
    ],
    [700, 2300, 900, 4460, 1000],
    font_size=7.7,
)
add_heading("17.2 External and Project Dependencies", 2)
add_bullets([
    "Confirmed business owner, prioritization authority, UAT representatives, and sample/source data.",
    "Valid Metronic license and approved component/plugin list.",
    "Hosting, TLS, database, queue/cache, email, object storage, monitoring, backup, and secret-management services.",
    "Decisions on authentication/MFA, tenant model, notification provider, FX source, data region, retention, and privacy terms.",
    "Approved calculation examples for representative vendors and on-demand/token rate structures.",
])
add_heading("17.3 Open Decisions Required Before Build Completion", 2)
add_table(
    ["Decision ID", "Decision", "Recommended baseline", "Needed by"],
    [
        ("D-01", "Single-user versus multi-user workspace at launch", "Design tenant-ready; enable owner plus optional viewer/editor roles.", "Architecture sign-off"),
        ("D-02", "Base currency and FX source", "User base currency with stored dated manual/approved provider rates.", "Cost module build"),
        ("D-03", "Tax display and recognition", "Store tax components; offer inclusive/exclusive reporting; no tax filing.", "Ledger design"),
        ("D-04", "Forecast default for on-demand", "Run-rate when data is fresh; manual/target fallback with visible method.", "Forecast build"),
        ("D-05", "Notification channels", "In-app and email for launch; add channels only after separate approval.", "Alert build"),
        ("D-06", "Financial-period close", "Optional monthly close with authorized reopen and adjustment workflow.", "Reporting build"),
        ("D-07", "Retention and hosting region", "Confirm by jurisdiction and service policy before production.", "Production readiness"),
        ("D-08", "MFA/step-up actions", "Offer MFA; require step-up for role changes, full export and account closure if feasible.", "Security sign-off"),
    ],
    [1100, 3000, 4160, 1100],
    font_size=7.9,
)
add_heading("17.4 Acceptance and Sign-off", 2)
add_p(
    "Approval of this BRD confirms the business problem, scope, requirements, rules, use cases, UI expectations, data/security baseline, "
    "test approach, and recommended delivery plan. Approval does not remove the need for detailed solution design, backlog refinement, "
    "data-protection/legal review, calculation examples, interface specifications, and sprint-level acceptance criteria."
)
add_table(
    ["Role", "Name", "Approval / conditions", "Signature", "Date"],
    [
        ("Business Sponsor", "", "", "", ""),
        ("Product Owner", "", "", "", ""),
        ("Technical Lead", "", "", "", ""),
        ("QA / Security Lead", "", "", "", ""),
    ],
    [1800, 1600, 3000, 1600, 1360],
    font_size=8.4,
)

page_break()

# Appendices
add_heading("Appendix A — Role and Permission Matrix", 1)
add_table(
    ["Capability", "Owner", "Admin", "Editor", "Viewer / Auditor", "Finance Reviewer"],
    [
        ("View dashboard/reports", "Yes", "Yes", "Yes", "Yes", "Yes"),
        ("Create/edit tools/subscriptions", "Yes", "Yes", "Yes", "No", "Read"),
        ("Change/cancel plan", "Yes", "Yes", "Yes*", "No", "Read"),
        ("Post ordinary cost/usage", "Yes", "Yes", "Yes", "No", "Yes"),
        ("Correct posted/closed cost", "Yes*", "Yes", "No", "No", "Yes"),
        ("Manage payment aliases", "Yes", "Yes*", "Yes*", "Masked read", "Masked read"),
        ("Manage targets/alerts", "Yes", "Yes", "Yes", "Read", "Read"),
        ("Manage users/roles/reference data", "Owner", "Yes", "No", "No", "No"),
        ("Export data", "Yes", "Yes", "Scoped", "Scoped", "Scoped"),
        ("Close account / purge", "Yes", "No", "No", "No", "No"),
        ("View audit", "Yes", "Yes", "Scoped", "Yes", "Yes"),
    ],
    [2500, 1000, 1000, 1100, 1900, 1860],
    font_size=7.9,
)
add_p(
    "* Permission is subject to workspace policy and may require step-up verification, impact preview, reason, or approval. "
    "The final matrix must be confirmed during solution design.",
    italic=True,
    color=MID_GRAY,
)

add_heading("Appendix B — Lifecycle Status Models", 1)
add_table(
    ["Object", "Statuses", "Key transition controls"],
    [
        ("Subscription", "Draft, Trial, Active, Paused, Pending Cancellation, Cancelled, Expired, Archived", "Effective dates, commercial confirmation, owner permission, reason and history."),
        ("Plan version", "Draft, Scheduled, Effective, Superseded, Cancelled", "No invalid overlap; immutable once historically used except correction workflow."),
        ("Cost entry", "Draft, Posted, Reconciled, Voided, Reversed, Refunded", "Posted entries corrected by linked reversal/replacement; closed periods restricted."),
        ("Usage entry", "Staged, Validated, Posted, Rejected, Superseded", "Provider/batch idempotency; allocation occurs only for eligible posted usage."),
        ("Payment account", "Active, Expiring, Inactive, Replaced, Archived", "Impact/reassignment required before future obligations use inactive account."),
        ("Target", "Draft, Active, Paused, Completed, Archived", "Evaluation only when active and period/scope valid."),
        ("Alert", "Open, Acknowledged, Snoozed, Assigned, Resolved, Dismissed, Auto-recovered", "Severity/state history retained; critical dismissal requires reason."),
        ("Import batch", "Uploaded, Staged, Validated, Processing, Completed, Partial, Failed, Rolled Back", "Idempotency, counts and row outcomes retained."),
    ],
    [1600, 3500, 4260],
    font_size=8.0,
)

add_heading("Appendix C — Sample Cost Calculations", 1)
add_heading("C.1 Annual Subscription Normalization", 2)
add_p(
    "Annual gross fee USD 144; discount USD 24; non-refundable tax USD 6. Net cash cost = 144 − 24 + 6 = USD 126. "
    "Normalized monthly management cost = 126 / 12 = USD 10.50. The cash view posts USD 126 at renewal; the normalized view shows USD 10.50 per month."
)
add_heading("C.2 On-Demand Tier Example", 2)
add_p(
    "Usage is 7,000,000 units. First 1,000,000 are included, next 4,000,000 cost USD 0.50 per million, and the remainder costs "
    "USD 0.40 per million. Cost = (4 × 0.50) + (2 × 0.40) = USD 2.80, plus any effective fixed fee/tax. "
    "The calculation explanation must show included and paid quantities separately."
)
add_heading("C.3 Target Forecast Example", 2)
add_p(
    "Monthly target USD 100; actual through day 15 of a 30-day period is USD 45; remaining committed cost is USD 30. "
    "Run-rate variable forecast is USD 45 / 15 × 30 = USD 90 only when the entire actual is treated as variable. "
    "If USD 30 of actual is recurring and USD 15 is variable, the preferred component forecast is USD 30 recurring actual + "
    "(USD 15 / 15 × 30) variable + USD 30 remaining commitments = USD 90. The UI must state which method/components were used."
)
add_heading("C.4 Multi-Currency Entry", 2)
add_p(
    "A charge of EUR 20 is posted using stored rate 1 EUR = 1.10 USD. Base amount is USD 22.00. Later FX changes do not rewrite "
    "the posted USD 22.00. A future commitment may use a newer approved rate and must show its forecast nature and rate date."
)

add_heading("Appendix D — Glossary and Abbreviations", 1)
add_table(
    ["Term", "Meaning"],
    [
        ("Actual", "Eligible posted net cost for the selected period."),
        ("Committed", "Known future obligation expected in the selected period."),
        ("Forecast", "Projected final period cost based on actuals, commitments and variable-cost assumptions."),
        ("MTD / EOM", "Month to date / end of month."),
        ("FX", "Foreign exchange conversion between original and reporting currency."),
        ("FIFO-by-expiry", "Consume the eligible balance that expires earliest first."),
        ("MoSCoW", "Must, Should, Could, Won't/Out-of-scope prioritization."),
        ("RPO / RTO", "Recovery point objective / recovery time objective."),
        ("WCAG", "Web Content Accessibility Guidelines."),
        ("PAN / CVV", "Card account number / card verification value; prohibited from storage in this product."),
        ("Idempotent", "Safe to retry without creating duplicate business effects."),
        ("Metronic", "Licensed UI template/component package used as the presentation foundation."),
    ],
    [1900, 7460],
    font_size=8.3,
)

add_heading("Appendix E — Definition of Ready and Definition of Done", 1)
add_heading("Definition of Ready", 2)
add_bullets([
    "Requirement/use case and priority are identified; business rule and calculation examples are clear.",
    "UX state, role/permission, validation, data fields, audit, notification and error behaviour are specified.",
    "Dependencies, migration/integration needs, security/privacy considerations and acceptance tests are understood.",
])
add_heading("Definition of Done", 2)
add_bullets([
    "Code, migration, authorization, validation, audit, responsive UI and accessibility behaviour are implemented and reviewed.",
    "Automated unit/feature/E2E tests pass; performance/security checks appropriate to risk are complete.",
    "Documentation, telemetry, operational runbook, migration/rollback and user guidance are updated.",
    "Product Owner accepts the requirement against representative data with no disqualifying defect.",
])

# Final document settings and metadata.
core = doc.core_properties
core.title = "Personal AI Subscription Management Platform — Business Requirements Document"
core.subject = "Business, functional, UX, data, technical and acceptance requirements"
core.author = "Product and Business Analysis"
core.keywords = "BRD, AI subscriptions, token usage, cost tracking, Laravel, MySQL, Metronic"
core.comments = "Version 1.0 baseline"

settings = doc.settings.element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

# Keep headings with the next paragraph and avoid widows/orphans globally.
for paragraph in doc.paragraphs:
    p_pr = paragraph._p.get_or_add_pPr()
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

# 14 Testing
add_heading("14. Testing Strategy and Acceptance Test Scenarios", 1)
add_heading("14.1 Test Strategy", 2)
add_bullets([
    "Unit tests cover money, recurrence, proration, tier pricing, allocation, FX, forecast, target, state transition, and permission rules.",
    "Feature/integration tests cover authenticated workflows, validation, database transactions, queues, schedules, imports, exports, and notification delivery.",
    "End-to-end tests cover the highest-value use cases on desktop and mobile breakpoints using realistic data and role combinations.",
    "Security tests cover authentication, authorization, tenant isolation, injection/XSS/CSRF, IDOR, rate limits, uploads, secret leakage, masking, and audit.",
    "Accessibility tests combine automated scanning with keyboard, focus, zoom/reflow, screen-reader spot checks, contrast, chart alternatives, and error handling.",
    "Performance tests use representative high-volume workspaces for dashboard, list, report, import, target evaluation, and queue backlog behaviour.",
    "Recovery tests validate backup restoration, failed-job retry, import idempotency, connector outage, email outage, and rollback/roll-forward.",
    "User acceptance testing is passed only when all Must requirements and critical journeys meet agreed results with no open critical/high defects.",
])
add_heading("14.2 Functional and Calculation Acceptance Tests", 2)
functional_tests = [
    ("AT-001", "UC-01 / Setup", "Complete onboarding with BDT base currency and Asia/Dhaka time zone; add masked card and monthly subscription.", "Dashboard reflects preferences, normalized cost, next renewal and starter target; no full card data accepted.", "E2E"),
    ("AT-002", "UC-02 / Monthly", "Create USD 20 monthly plan starting on the 15th with 10% discount and fixed tax.", "Schedule, first-period proration rule, net cost, next charge and annualized view match configured rules.", "Calc"),
    ("AT-003", "UC-02 / Annual", "Create annual plan USD 120 renewing in six months.", "Cash forecast shows renewal in correct month; normalized monthly view shows USD 10 before other components.", "Calc"),
    ("AT-004", "Recurring dates", "Create plan billed on day 31 and generate commitments across February/leap year.", "End-of-month rule is applied consistently with no skipped or duplicate commitment.", "Calc"),
    ("AT-005", "UC-03 / Tier rate", "Configure first 1M units included, next 4M at rate A, remainder at rate B; post 7M.", "Included units and tiers apply once; total and explanation match hand calculation.", "Calc"),
    ("AT-006", "UC-03 / Hybrid", "Configure base fee with included units and overage; add usage below then above allowance.", "Forecast/actual show base only below allowance and correct overage above it.", "Calc"),
    ("AT-007", "UC-04 / Package", "Buy two packages with different expiry; post consumption.", "FIFO-by-expiry allocation, remaining balances and effective paid-unit cost reconcile.", "Calc"),
    ("AT-008", "Package expiry", "Let a package expire with unused units.", "Balance is no longer spendable; expired quantity/cost impact remains visible and alert resolves appropriately.", "Func"),
    ("AT-009", "UC-05 / Plan change", "Future-date price increase and payment account change.", "Old schedule/history retained; new commitments and account apply from effective date only.", "E2E"),
    ("AT-010", "Version overlap", "Attempt overlapping effective plan/rate ranges.", "Activation is blocked with field-level explanation; no partial version is created.", "Neg"),
    ("AT-011", "UC-06 / Reconcile", "Post actual charge matching one expected commitment.", "Expected item is matched/offset exactly once; actual and forecast totals do not double count.", "E2E"),
    ("AT-012", "Duplicate charge", "Import same provider reference twice.", "Second row is skipped/flagged according to decision; ledger remains idempotent.", "Neg"),
    ("AT-013", "Refund", "Post refund linked to original charge.", "Net cost, target utilization and audit chain update; original remains visible.", "Func"),
    ("AT-014", "UC-13 / Correction", "Correct a posted cost by reversal and replacement.", "No overwrite occurs; chain balances and reports use corrected net result.", "E2E"),
    ("AT-015", "FX history", "Post foreign charge with stored rate; later update current FX rate.", "Historical base amount remains unchanged; new forecasts use the approved effective rate.", "Calc"),
]
add_test_table(functional_tests)

add_heading("14.3 Target, Alert, Renewal and Payment Acceptance Tests", 2)
alert_tests = [
    ("AT-016", "UC-07 / Target", "Create monthly category target with warning 80% and critical 100% using forecast basis.", "Utilization and preview are correct; invalid reversed thresholds are rejected.", "E2E"),
    ("AT-017", "UC-08 / Warning", "Post charge causing forecast to cross 80%.", "One warning alert is created and delivered with contributors and correct action link.", "E2E"),
    ("AT-018", "Alert dedupe", "Recalculate same unchanged warning multiple times within cool-down.", "No duplicate user notification; evaluation/delivery trace remains inspectable.", "Func"),
    ("AT-019", "Alert escalation", "Increase cost from warning to critical during cool-down.", "Existing event escalates and critical notification is sent immediately per policy.", "Func"),
    ("AT-020", "Alert recovery", "Reverse erroneous cost so utilization drops below threshold.", "Event records auto-recovery/updated state but prior warning history is retained.", "Func"),
    ("AT-021", "Quiet hours", "Trigger non-critical alert during quiet hours.", "In-app event exists; configured outbound delivery waits and sends after quiet period.", "Func"),
    ("AT-022", "Delivery failure", "Simulate email provider failure.", "Delivery retries with backoff, records failure, leaves in-app alert open, and surfaces operations issue.", "Integr"),
    ("AT-023", "Renewal reminders", "Active auto-renew plan enters 30/14/7/1-day windows.", "Configured reminders occur once per window in workspace time zone with correct amount and deadline.", "Sched"),
    ("AT-024", "Trial conversion", "Trial ending with automatic conversion enabled.", "Reminder shows conversion price; forecast includes post-trial fee according to configured rule.", "Calc"),
    ("AT-025", "Cancellation deadline", "Cancel before deadline and confirm effective/access-end dates.", "Eligible future commitments stop; historical costs and evidence remain.", "E2E"),
    ("AT-026", "Late cancellation", "Record cancellation after contractual deadline.", "System warns renewal may remain, requires outcome confirmation, and does not remove commitment prematurely.", "Neg"),
    ("AT-027", "UC-10 / Replace", "Bulk replace expiring card for selected subscriptions.", "Future plan versions change; past entries retain old alias; excluded items are reported.", "E2E"),
    ("AT-028", "Account expiry", "Card expiry precedes next renewal.", "Alert lists every affected accessible subscription and next charge without exposing sensitive data.", "Func"),
    ("AT-029", "Payment deactivation", "Deactivate account still linked to future commitments.", "System blocks or requires guided reassignment/exception; impact count is shown.", "Neg"),
    ("AT-030", "On-demand run rate", "Post half-month usage and compute run-rate forecast.", "Elapsed/total billable-day rule, minimum/cap and data freshness are displayed and correct.", "Calc"),
]
add_test_table(alert_tests)

add_heading("14.4 Import, Reporting, UI and Responsive Acceptance Tests", 2)
ui_tests = [
    ("AT-031", "UC-11 / Import", "Upload file containing valid, warning, error and duplicate rows.", "Preview counts and row messages are accurate; execution follows documented partial/atomic policy.", "E2E"),
    ("AT-032", "Import retry", "Force worker failure after batch staging and retry.", "Committed rows are not duplicated; batch resumes or rolls back deterministically.", "Recovery"),
    ("AT-033", "UC-12 / Dashboard", "Apply period/category/account filters and drill from forecast KPI.", "All cards/charts use same context; underlying rows reconcile to KPI and URL/state preserves filters.", "E2E"),
    ("AT-034", "Empty state", "Open new blank workspace.", "No broken zero charts; onboarding guidance and permitted quick actions are shown.", "UI"),
    ("AT-035", "Stale source", "Vendor usage has not refreshed by expected time.", "Dashboard/report marks stale data and explains forecast inclusion/exclusion.", "UI"),
    ("AT-036", "Export masking", "Viewer exports payment-account report.", "Export matches filters and contains only permitted masked fields; event is audited.", "Security"),
    ("AT-037", "Desktop layout", "Test core pages at 1440×900.", "Navigation, KPI/cards, forms and tables align; no overlap, clipping or unexplained horizontal page scroll.", "Visual"),
    ("AT-038", "Tablet layout", "Test core journeys at 768×1024 and landscape.", "Drawer navigation, two-column cards and priority-table behaviour remain complete and usable.", "Visual"),
    ("AT-039", "Mobile layout", "Complete add subscription, record usage and resolve alert at 375×812.", "Single-column flow, keyboard/input types, sticky actions and validation work without overlap or page-wide scrolling.", "E2E"),
    ("AT-040", "Small mobile", "Test dashboard and forms at 320 px and 200% zoom.", "Core content reflows; no action is hidden; exceptions are limited to genuine tables with accessible alternative.", "A11y"),
    ("AT-041", "Keyboard", "Complete core flows without mouse.", "Focus order is logical and visible; menus/modals/drawers work; focus returns correctly.", "A11y"),
    ("AT-042", "Screen reader", "Review dashboard KPIs, alert detail, form errors and chart alternative.", "Names, roles, states, error associations and summaries are understandable.", "A11y"),
    ("AT-043", "Color independence", "Inspect target/alert states under grayscale and color-vision simulation.", "Severity remains clear through labels, icons and text; contrast passes target.", "A11y"),
    ("AT-044", "Validation persistence", "Submit invalid conditional form and receive server errors.", "Specific errors appear at fields and summary; valid inputs remain; focus moves to summary/first error.", "UI"),
    ("AT-045", "Saved view", "Save filters/columns/density, reload on mobile and desktop.", "Allowed preferences persist; incompatible mobile columns degrade cleanly and can be reset.", "Func"),
]
add_test_table(ui_tests)

add_heading("14.5 Security, Performance and Operational Acceptance Tests", 2)
ops_tests = [
    ("AT-046", "Tenant isolation", "Attempt direct-object and crafted-query access to another workspace across web, API, export and job paths.", "Every attempt is denied without existence leakage; security event is logged appropriately.", "Security"),
    ("AT-047", "Role matrix", "Exercise create/edit/correct/export/admin actions as each role.", "Server-side permission outcomes match approved matrix; hidden buttons are not the only control.", "Security"),
    ("AT-048", "Sensitive input", "Attempt to enter full card number, CVV or password in payment fields/notes.", "Dedicated fields never request them; prohibited data is rejected/warned and not logged.", "Security"),
    ("AT-049", "Upload safety", "Upload oversized, mismatched MIME, executable and infected test files.", "Unsafe files are rejected/quarantined; no public or executable access is created.", "Security"),
    ("AT-050", "Session security", "Test reset token replay, session fixation, inactivity, logout all devices and rate limiting.", "Tokens are single-use/expiring; sessions rotate/revoke; abuse receives safe responses.", "Security"),
    ("AT-051", "Audit integrity", "Change plan, correct cost, change role, export data and resolve alert.", "Complete append-only audit events include actor, outcome, entity, correlation and required reason.", "Audit"),
    ("AT-052", "Dashboard performance", "Load representative high-volume workspace under agreed concurrent traffic.", "95th percentile meets target; totals are correct; cache does not cross workspace/filter boundaries.", "Perf"),
    ("AT-053", "Queue backlog", "Create high alert/import workload and restart workers.", "Jobs are idempotent, recover, expose progress/failure and do not lose/duplicate financial events.", "Recovery"),
    ("AT-054", "Backup restore", "Restore latest approved backup into isolated environment and run reconciliation checks.", "Restore meets RPO/RTO baseline and source ledger totals/snapshots reconcile.", "Recovery"),
    ("AT-055", "Scheduler/time zone", "Run renewal and target jobs across DST/time-zone boundary and month-end.", "Events use workspace commercial dates, execute once, and do not drift with server time.", "Sched"),
    ("AT-056", "Connector outage", "Simulate timeout, partial payload and duplicate replay.", "Timeout/circuit/retry rules apply; data is idempotent; stale status and manual fallback are visible.", "Integr"),
    ("AT-057", "Metronic hygiene", "Inspect production bundle/routes/assets.", "Unused demos/plugins/credentials are absent; dependency/license inventory is documented.", "Security"),
    ("AT-058", "Error privacy", "Trigger validation, 404/403/500 and failed background jobs.", "User receives safe correlation code; logs contain diagnostic context without secrets or prohibited data.", "Security"),
    ("AT-059", "Account closure", "Export then close account with retention hold.", "Sessions revoke, notifications stop, retained classes are explained, purge workflow respects hold.", "Privacy"),
    ("AT-060", "Period close", "Close month, attempt ordinary edit, then perform authorized reopen with reason.", "Edit is blocked while closed; reopen and affected recalculation are fully audited.", "Func"),
]
add_test_table(ops_tests)

add_heading("14.6 Release Acceptance Gates", 2)
add_bullets([
    "All Must requirements are implemented or have an approved, documented waiver from the Product Owner and Sponsor.",
    "All UC-01 through UC-15 primary flows pass; no Critical/High unresolved defect affects security, money, permission, data loss, renewal, or alert accuracy.",
    "Cost/forecast reference calculations reconcile at defined decimal precision and rounding tolerance.",
    "Responsive and accessibility acceptance passes for the supported browser/device matrix and core workflows.",
    "Security review, dependency scan, backup restore, migration rehearsal, monitoring, support runbooks, and production-readiness checklist are complete.",
    "Business owner signs off dashboard totals, target/alert behaviour, imports, exports, privacy handling, and notification templates using representative data.",
])

# 15 traceability
add_heading("15. Requirements Traceability", 1)
trace_rows = [
    ("Identity & setup", "OBJ-01, OBJ-06", "FR-001–006, FR-083", "UC-01, UC-15", "AT-001, AT-034, AT-050, AT-059"),
    ("Catalogue & types", "OBJ-01", "FR-010–015, FR-080", "UC-02, UC-03, UC-14", "AT-002–005, AT-010"),
    ("Recurring lifecycle", "OBJ-01, OBJ-03", "FR-020–028", "UC-02, UC-05, UC-09", "AT-002–004, AT-009–010, AT-023–026"),
    ("Cost & reconciliation", "OBJ-02, OBJ-04", "FR-030–037", "UC-06, UC-11, UC-13", "AT-011–015, AT-031–032, AT-060"),
    ("Token & on-demand", "OBJ-02, OBJ-05", "FR-040–048", "UC-03, UC-04", "AT-005–008, AT-030"),
    ("Payment accounts", "OBJ-04", "FR-050–055", "UC-02, UC-10", "AT-001, AT-027–029, AT-036, AT-048"),
    ("Targets & alerts", "OBJ-02, OBJ-03", "FR-060–068", "UC-07, UC-08", "AT-016–022"),
    ("Dashboard & reports", "OBJ-02, OBJ-05, OBJ-06", "FR-070–077", "UC-12", "AT-033–045, AT-052"),
    ("Admin & operations", "OBJ-01–06", "FR-080–087", "UC-11, UC-14, UC-15", "AT-031–032, AT-046–060"),
    ("Security & privacy", "OBJ-04", "SEC-01–12, NFR-06–15", "All", "AT-036, AT-046–059"),
]
add_table(
    ["Capability", "Objectives", "Requirements", "Use cases", "Acceptance tests"],
    trace_rows,
    [1500, 1200, 2400, 1700, 2560],
    font_size=7.6,
)

page_break()

# 8 UX and responsive UI
add_heading("8. User Experience and Responsive UI Requirements", 1)
add_heading("8.1 UX Principles", 2)
add_bullets([
    "Compact, not cramped: prioritize current cost, next risk, and next action; reveal detail progressively.",
    "One calculation language: Actual, Committed, Forecast, Target, and Variance use consistent labels and colors everywhere.",
    "Explain every number: totals link to contributing records, formula basis, data freshness, currency conversion, and exclusions.",
    "Fast entry: sensible defaults, conditional fields, duplicate/quick-add, keyboard support, inline validation, and saved mappings.",
    "Safe actions: preview financial impact, confirm destructive lifecycle actions, support undo where safe, and retain history.",
    "Responsive by design: the same capability is available on mobile, but layout adapts to touch, narrow width, and short sessions.",
    "Accessible and calm: color is never the only signal; alerts are prioritized; empty, loading, stale, error, and offline states are explicit.",
])
add_heading("8.2 Information Architecture", 2)
add_table(
    ["Primary navigation", "Key destinations"],
    [
        ("Dashboard", "Overview, targets, alerts, spend trend, renewals, usage, payment risks, quick actions."),
        ("Subscriptions", "All subscriptions, tools, vendors, calendar, trials, renewals, archived records."),
        ("Usage & Tokens", "Meters, usage entries, packages/balances, imports, models/projects, forecasts."),
        ("Costs", "Ledger, expected commitments, reconciliation, adjustments, closed periods, imports."),
        ("Targets & Alerts", "Targets, policies, alert inbox, notification history, recommendations."),
        ("Payment Accounts", "Cards/accounts/wallets, linked subscriptions, expiry, limits, reassignment."),
        ("Reports", "Cost summary, variance, renewal, utilization, payment source, token, savings, exports."),
        ("Settings", "Profile, workspace, categories, subscription types, roles, imports, integrations, audit, operations."),
    ],
    [2000, 7360],
    font_size=8.5,
)
add_heading("8.3 Screen Catalogue and UI Acceptance", 2)
add_table(
    ["Screen", "Essential content", "Primary actions", "Responsive behaviour"],
    [
        ("Sign in / Recovery", "Email, password/SSO, verification and safe errors", "Sign in, reset, verify", "Single column; large touch targets; password manager compatible."),
        ("Onboarding", "Progress, preferences, payment alias, first subscription, target", "Back, skip, save, finish", "Step layout becomes stacked; sticky bottom action bar."),
        ("Dashboard", "KPI strip, target bars, alerts, renewals, charts, quick actions", "Filter, drill, add, export", "Cards stack 1-up; charts have compact/table view; filters in sheet."),
        ("Subscription List", "Name, type, status, owner, cost, renewal, payment alias, alert", "Search, filter, sort, bulk action", "Desktop table switches to compact cards or priority columns."),
        ("Subscription Detail", "Overview, plan history, costs, usage, targets, alerts, attachments, audit", "Edit, change plan, pause, cancel", "Tabs become horizontally scrollable or select menu; actions in bottom sheet."),
        ("Add/Edit Subscription", "Conditional form by type with live cost preview", "Save draft, validate, activate", "One-column sections; numeric keypad; dates and selects touch friendly."),
        ("Usage & Token Detail", "Balances, expiry, rate tiers, consumption, cost, forecast", "Add/import usage, buy package, adjust", "Summary fixed above collapsible detailed tables."),
        ("Cost Ledger", "Posted/draft/expected entries, matching and variance", "Add, import, reconcile, correct", "Priority columns plus row detail drawer; bulk tools in overflow."),
        ("Target Editor", "Scope, period, amount, basis, thresholds, channels, preview", "Test, activate, duplicate", "Form stacks; preview stays visible as compact summary."),
        ("Alert Inbox / Detail", "Severity, state, variance, contributors, delivery, actions", "Acknowledge, assign, snooze, resolve", "List cards; action buttons remain reachable by thumb."),
        ("Payment Account Detail", "Masked identity, expiry, spend, linked plans, next charges", "Edit, replace, deactivate", "No sensitive data exposed; linked items rendered as cards."),
        ("Reports", "Report selector, filters, summary, chart/table, freshness", "Run, save, export, schedule", "Filters in sheet; printable rendering remains desktop-width."),
        ("Settings / Admin", "Reference data, roles, alerts, audit, operations", "Create, edit, archive, test", "Grouped navigation; avoid very wide editable grids."),
    ],
    [1500, 3200, 1900, 2760],
    font_size=7.5,
)
add_heading("8.4 Dashboard Composition", 2)
add_bullets([
    "Top bar: workspace, global search, period selector, currency view, notifications, profile, and Add button.",
    "KPI row: Actual MTD, Committed Remaining, Forecast EOM, Target Variance, Next 30-Day Renewals, and Active Critical Alerts.",
    "Priority area: target progress, critical/warning alert list, renewal/trial timeline, and payment-expiry conflicts.",
    "Insight area: cost trend, category/tool mix, recurring versus on-demand, token cost/usage, and top increases.",
    "Action area: add subscription, record cost, record usage, buy package, create target, import, and run reconciliation.",
    "Every tile shows period, currency, data freshness, accessible label, and drill-down destination; ambiguous standalone percentages are prohibited.",
])
add_heading("8.5 Responsive Breakpoints and Layout", 2)
add_table(
    ["Viewport", "Layout requirements"],
    [
        ("≥1200 px", "Persistent sidebar; 12-column grid; dense table view; multi-panel detail where useful."),
        ("992–1199 px", "Collapsible sidebar; 8/12-column grid; reduced secondary columns; no horizontal page scroll."),
        ("768–991 px", "Drawer navigation; 2-column cards; tables use priority columns and row details."),
        ("≤767 px", "Single-column content; bottom/sticky primary actions; filters in full-height sheet; cards replace wide tables."),
        ("≤375 px", "Core add/edit/alert tasks remain operable; labels wrap; touch targets do not overlap; charts offer text/table alternative."),
    ],
    [1500, 7860],
    font_size=8.4,
)
add_heading("8.6 Form, Table and State Requirements", 2)
add_bullets([
    "Required fields are identified in label text; server and client validation use the same rule and specific corrective messages.",
    "Conditional fields appear only for the selected subscription/cost model, while changing type warns before discarding entered values.",
    "Money inputs show currency, accept appropriate decimal precision, preserve unrounded value, and never use floating-point arithmetic server-side.",
    "Dates show time-zone context when timing matters; renewal, cancellation deadline, and access-end date remain distinct.",
    "Tables have sticky headers on desktop, column chooser, sort, filter, pagination or virtualization, accessible row actions, and export matching filters.",
    "All data views define loading skeleton, no-data, no-results, partial/stale data, permission denied, recoverable error, and fatal error states.",
    "Bulk actions show selected count and an impact preview; sensitive or destructive bulk changes require explicit confirmation.",
    "Success notices name the saved object and provide a useful next action; errors preserve entered data.",
])
add_heading("8.7 Accessibility and Usability", 2)
add_bullets([
    "Meet WCAG 2.2 AA as the product target, subject to formal accessibility validation before release.",
    "All functions operate by keyboard with visible focus, logical order, skip links, and no keyboard traps.",
    "Touch targets are at least 44×44 CSS pixels where practical; compact density never reduces critical action targets below accessible size.",
    "Text and meaningful UI components meet contrast requirements; warning/critical states use icon, label, and text in addition to color.",
    "Charts have descriptive titles, accessible summaries, tooltips reachable by keyboard, and equivalent data tables.",
    "Modals, drawers, menus, tabs, toasts, and validation messages expose correct accessible names, roles, state, and focus restoration.",
    "The interface supports 200% text zoom and browser zoom without loss of core functionality or two-dimensional scrolling, except genuine data tables.",
])
add_heading("8.8 Metronic Application Guidance", 2)
add_p(
    "Use Metronic as the visual/component foundation, not as the source of business rules. Establish an application design layer that maps "
    "Metronic tokens and components to approved semantic roles: primary action, neutral action, success, warning, critical, disabled, and focus. "
    "Remove unused demo assets, pages, plugins, and example credentials. Wrap vendor components so accessibility, validation, authorization, "
    "responsive behaviour, and future template upgrades remain controllable by the application."
)

# 9 Data requirements
add_heading("9. Data Requirements and Logical Data Model", 1)
add_heading("9.1 Logical Entity Catalogue", 2)
entity_rows = [
    ("Workspace", "Base currency, locale, time zone, fiscal settings, feature and retention policy.", "1:N users, tools, targets"),
    ("User / Role", "Identity, status, preferences, role memberships.", "N:M workspace roles"),
    ("Vendor", "Supplier identity, URL, support and billing metadata.", "1:N tools"),
    ("Tool", "Product identity, category, purpose, owner, tags, status.", "1:N subscriptions"),
    ("Category / Tag", "Configurable classification and presentation metadata.", "N:M tools/subscriptions"),
    ("Subscription Type", "Commercial model and enabled characteristics.", "1:N subscriptions"),
    ("Subscription", "Lifecycle, owner, plan, status, dates, vendor account reference.", "1:N plan versions"),
    ("Plan Version", "Effective terms, billing cadence, amount, tax/discount, auto-renew, payment account.", "1:N cost components/rates"),
    ("Cost Component", "Base fee, seat, add-on, usage, tax, surcharge, discount semantics.", "1:N schedules/entries"),
    ("Rate / Tier", "Meter unit, range, price, minimum/cap, included units, effective date.", "N:1 plan version"),
    ("Usage Entry", "Period, project/model, units, source, provider reference, status.", "N:1 meter/subscription"),
    ("Token Package / Lot", "Purchased/granted units, cost, currency, expiry and balance.", "1:N allocations"),
    ("Usage Allocation", "Units allocated from a package/included/overage source.", "N:1 usage and lot"),
    ("Cost Entry", "Ledger event, original/base amounts, dates, status, references.", "N:1 subscription/account"),
    ("Expected Commitment", "Scheduled future obligation and matching state.", "N:1 plan component"),
    ("Payment Account", "Masked payment source, owner, currency, expiry, status.", "1:N plans/entries"),
    ("Target", "Scope, period, amount, basis, thresholds, status.", "1:N evaluations"),
    ("Alert Policy", "Event type, scope, thresholds, channels, recipients and schedule.", "1:N alerts"),
    ("Alert / Delivery", "Event state, severity, message, deduplication, acknowledgement and channel outcome.", "N:1 policy"),
    ("Attachment", "Safe file metadata, storage key, hash, classification.", "Polymorphic authorized link"),
    ("Import Batch / Row", "Mapping, file, counts, validation and processing outcome.", "1:N staged rows"),
    ("Audit Event", "Actor, action, entity, correlation, outcome, before/after summary.", "Append-only"),
]
add_table(
    ["Entity", "Purpose / key data", "Primary relationship"],
    entity_rows,
    [2000, 5400, 1960],
    font_size=7.7,
)
add_heading("9.2 Data Integrity Requirements", 2)
add_bullets([
    "Use surrogate primary keys and workspace/tenant keys on all owned records; use unique constraints for business identifiers within scope.",
    "Use DECIMAL for money, FX, rates, and high-precision usage; never use FLOAT/DOUBLE for financial values.",
    "Store date-only commercial deadlines separately from UTC timestamps; store workspace time zone used for interpretation.",
    "Plan/rate effective ranges must not overlap for the same commercial component unless explicitly modeled as additive.",
    "Ledger totals derive from immutable posted entries and linked reversals; cached summaries are rebuildable from source transactions.",
    "Foreign keys, state constraints, idempotency keys, provider references, and unique batch-row keys prevent orphaning and duplication.",
    "Every imported or automated record retains source, batch/connector, provider reference, received time, and original payload hash/trace as permitted.",
    "Soft deletion is used for records with historical references; purge is a separate permissioned retention operation.",
])
add_heading("9.3 Data Classification", 2)
add_table(
    ["Class", "Examples", "Handling"],
    [
        ("Public/Reference", "Tool/vendor public name, category labels", "Normal access within product context."),
        ("Internal", "Subscription plan, cost, target, notes, usage", "Authenticated, role-controlled, encrypted in transit/at rest."),
        ("Confidential", "Payment aliases, invoice attachments, exports, audit history", "Least privilege, masking, access logging, retention controls."),
        ("Restricted / prohibited", "Full PAN, CVV, bank password, vendor password, raw secret/API token in ordinary fields", "Do not collect. Approved connector secrets require dedicated encrypted secret storage and rotation."),
    ],
    [1800, 3200, 4360],
    font_size=8.3,
)
add_heading("9.4 Retention and Archival Baseline", 2)
add_p(
    "Retention periods are configurable by record class and jurisdiction. As a baseline, keep financial and audit history for the approved "
    "management-reporting period, keep import staging files only long enough to validate and recover the batch, and expire generated exports "
    "quickly. Purge must verify legal/financial holds, dependencies, backups, and workspace ownership before execution."
)

# 10 Architecture
add_heading("10. Solution Architecture and Technical Requirements", 1)
add_heading("10.1 Target Architecture", 2)
add_p(
    "The solution will be a modular Laravel web application with server-rendered and/or approved reactive UI components using Metronic, "
    "a MySQL transactional database, background queues, scheduled jobs, object storage for attachments/exports, cache/lock service, and "
    "integrations behind explicit adapters. A modular monolith is preferred for the initial release to preserve transactional consistency "
    "and operational simplicity while maintaining clear bounded modules."
)
add_table(
    ["Layer / module", "Responsibilities"],
    [
        ("Presentation", "Metronic layout and components, responsive views, form requests, accessible charts/tables, API/web endpoints."),
        ("Identity & Workspace", "Authentication, authorization, preferences, tenant/workspace boundary, session/security events."),
        ("Catalogue & Subscription", "Vendors, tools, types, plan versions, lifecycle, renewal schedules, attachments."),
        ("Cost & Reconciliation", "Ledger, expected commitments, imports, matching, FX, close/reopen and corrections."),
        ("Usage & Pricing", "Meters, rate versions, packages, allocations, usage costing and on-demand forecasting."),
        ("Budget & Alerting", "Targets, evaluations, forecasts, alert policy, deduplication, escalation and delivery."),
        ("Reporting", "Read models, filters, KPI calculations, exports, scheduled reports and data freshness."),
        ("Operations & Audit", "Jobs, health, integration runs, audit, retention, backups and administration."),
    ],
    [2200, 7160],
    font_size=8.3,
)
add_heading("10.2 Laravel Implementation Requirements", 2)
add_bullets([
    "Use Form Request or equivalent server-side validation, policies/gates for authorization, service/domain classes for calculations, and database transactions for multi-record financial operations.",
    "Use migrations, seeders for system reference data, factories for tests, and explicit indexes for workspace, status, dates, references, and report filters.",
    "Queue email, imports, exports, connectors, forecast recomputation, and heavy reports; use idempotent jobs with retry/backoff and dead-letter visibility.",
    "Use the scheduler for recurring commitments, renewal/trial/payment checks, target evaluation, data freshness, report delivery, retention, and health checks.",
    "Treat alerts and audit as domain events/outbox-backed where necessary so database changes and external delivery remain consistent.",
    "Centralize money, period, recurrence, FX, pricing-tier, and forecast logic in tested domain services; views must not duplicate formulas.",
    "Implement versioned application interfaces for integrations; use circuit breakers/timeouts and retain manual fallback.",
    "Use configuration/environment secrets outside source control; production debugging must not expose financial or credential data.",
])
add_heading("10.3 MySQL Requirements", 2)
add_bullets([
    "Use InnoDB, utf8mb4, explicit foreign keys, UTC timestamps, decimal columns sized for expected amounts/units, and consistent collation.",
    "Index by workspace plus common filters; verify composite indexes using representative queries and data volumes.",
    "Use transactions and row/application locks for balance allocation, duplicate-sensitive imports, reconciliation, plan changes, and target state transitions.",
    "Prefer normalized source-of-truth tables; use summary tables/materialized read models only when refresh and rebuild rules are explicit.",
    "Partitioning or archival tables may be introduced only after measured volume warrants them and recovery/query behaviour is tested.",
])
add_heading("10.4 Integration Boundaries", 2)
add_table(
    ["Integration", "Direction", "Baseline behaviour", "Failure behaviour"],
    [
        ("Email provider", "Outbound", "Alerts, verification, recovery, digests, reports", "Retry; in-app record remains; admin sees failure."),
        ("FX-rate provider", "Inbound", "Approved dated conversion rates", "Use last approved/manual rate with freshness warning; never invent."),
        ("Vendor usage/billing API", "Inbound", "Usage, charges, balances where authorized", "Mark stale, retry, and preserve manual/import path."),
        ("Object storage / malware scan", "Both", "Attachments and exports with signed access", "Reject unsafe file; no public bucket or permanent URL."),
        ("Optional bank/card feed", "Inbound", "Masked transactions for reconciliation", "Connector isolated; no payment initiation; unmatched data retained safely."),
        ("Optional notification channel", "Outbound", "Browser push/SMS/chat after approval", "Channel-specific failure logged without suppressing in-app alert."),
    ],
    [1700, 900, 3850, 2910],
    font_size=7.8,
)

# 11 Security
add_heading("11. Security, Privacy, Audit and Compliance", 1)
add_heading("11.1 Security Requirements", 2)
security_rows = [
    ("SEC-01", "Enforce authentication and server-side authorization for every data access and action."),
    ("SEC-02", "Separate workspace/tenant data in every query, cache key, job, export, attachment path, and audit search."),
    ("SEC-03", "Protect against CSRF, XSS, SQL injection, mass assignment, insecure direct-object reference, file upload abuse, and session fixation using framework controls and tests."),
    ("SEC-04", "Use TLS in transit and platform/database/storage encryption at rest; manage secrets in an approved secret store."),
    ("SEC-05", "Never store full card number, CVV, banking password, or vendor password; reject these patterns where feasible and educate users."),
    ("SEC-06", "Mask payment aliases in UI, notifications, logs, exports, screenshots, and support tools according to role."),
    ("SEC-07", "Rate-limit sign-in, recovery, imports, exports, alert tests, and connector endpoints; detect abuse without leaking account existence."),
    ("SEC-08", "Use secure cookies, session rotation, expiry, logout-all-devices, and optional MFA/step-up for exports, role changes, or account closure."),
    ("SEC-09", "Validate MIME/content, size, extension, malware result, ownership, and signed access for attachments and generated files."),
    ("SEC-10", "Audit sensitive reads/exports and all material writes; protect audit events from ordinary modification."),
    ("SEC-11", "Log application events with correlation IDs and redaction; logs must exclude secrets, raw sensitive payloads, and unnecessary financial detail."),
    ("SEC-12", "Patch supported dependencies, scan source/dependencies, review Metronic/plugins, and remediate vulnerabilities under an agreed severity SLA."),
]
add_table(["ID", "Requirement"], security_rows, [1000, 8360], font_size=8.2)
add_heading("11.2 Privacy Requirements", 2)
add_bullets([
    "Collect only data required for subscription management; explain purpose at entry and in privacy notice.",
    "Support user access/export, correction, deactivation, and deletion/purge where retention obligations permit.",
    "Separate consent/authorization for vendor, financial-feed, and external notification integrations.",
    "Define data processor/controller roles, hosting region, subprocessors, breach process, and jurisdiction-specific duties before launch.",
    "Use synthetic or anonymized data in non-production; production copies require approval and masking.",
])
add_heading("11.3 Audit Event Minimum Fields", 2)
add_p(
    "Event ID, workspace, UTC time, actor/service, role, action, outcome, entity type/ID, before/after summary or changed fields, "
    "reason where required, IP/device context as permitted, correlation/request ID, import/job/connector reference, and retention class."
)

# 12 NFR
add_heading("12. Non-Functional Requirements", 1)
nfr_rows = [
    ("NFR-01", "Availability", "Target 99.5% monthly for production excluding approved maintenance; final SLA confirmed before launch."),
    ("NFR-02", "Interactive performance", "95th percentile server response under 2 seconds for common list/detail actions at agreed load; dashboard under 3 seconds with cached summaries."),
    ("NFR-03", "Heavy work", "Imports, exports, connector sync, large reports and recalculation run asynchronously with progress and safe retry."),
    ("NFR-04", "Scale baseline", "Support at least 10,000 subscriptions, 1,000,000 ledger/usage rows per workspace, and 100 concurrent active sessions, subject to load-test refinement."),
    ("NFR-05", "Reliability", "Financial operations are atomic/idempotent; duplicate delivery or job retry must not duplicate cost or usage."),
    ("NFR-06", "Recovery", "Baseline RPO ≤24 hours and RTO ≤8 hours; production target refined by sponsor; restore drills at least twice yearly."),
    ("NFR-07", "Accessibility", "WCAG 2.2 AA target for supported core workflows, verified by automated and manual testing."),
    ("NFR-08", "Responsive compatibility", "Current major Chrome, Edge, Firefox, Safari and supported mobile Safari/Chrome; graceful behaviour from 320 px width."),
    ("NFR-09", "Maintainability", "Modular code, automated migrations, static analysis, tests, coding standard, documented calculations and public interfaces."),
    ("NFR-10", "Observability", "Health, structured logs, metrics, queue depth, failed jobs, alert delivery, import/connector outcomes and calculation latency."),
    ("NFR-11", "Localization", "Locale-aware date/number/currency formatting; user text supports Unicode; business calculations remain locale independent."),
    ("NFR-12", "Data portability", "Human-readable and machine-readable exports preserve IDs, dates, currency and relationship references where appropriate."),
    ("NFR-13", "Supportability", "User-facing correlation code, admin diagnostics, safe impersonation only if approved and audited, and runbooks for common failures."),
    ("NFR-14", "Deployment", "Separate environments, automated CI/CD gates, backward-compatible migrations, rollback/roll-forward plan, and seeded reference data."),
    ("NFR-15", "Time accuracy", "NTP-synchronized infrastructure; UTC storage plus explicit workspace time-zone conversion and DST tests."),
]
add_table(
    ["ID", "Attribute", "Requirement / target"],
    nfr_rows,
    [900, 1700, 6760],
    font_size=8.0,
)

# 13 Reporting
add_heading("13. Reporting, Analytics and Notification Catalogue", 1)
add_heading("13.1 Standard Reports", 2)
add_table(
    ["Report", "Core measures and dimensions"],
    [
        ("Executive Cost Summary", "Actual, committed, forecast, target, variance, recurring/on-demand, trend; by period, owner, category, vendor, tool."),
        ("Subscription Inventory", "Status, type, plan, normalized cost, owner, payment account, renewal, cancellation deadline, data freshness."),
        ("Renewal & Trial Calendar", "Renewal/trial/cancellation/payment dates, amount, risk, decision and owner."),
        ("Target Variance", "Target scope/period/basis, utilization, warning/critical crossing, contributing costs and actions."),
        ("Token & Usage", "Units, included/paid/free, package balances, expiry, effective unit cost, forecast and overage."),
        ("Payment Account Exposure", "Spend, commitments, linked subscriptions, next charges, currency, expiry and exceptions."),
        ("Reconciliation", "Expected vs posted, matched/unmatched, variance, duplicate candidates, stale imports and closed-period status."),
        ("Savings Opportunities", "Unused/low-use, duplicates, price increase, cancellation window, projected avoidable cost and confidence."),
        ("Audit & Change", "Actor, action, outcome, entity, changed fields, reason, correlation and export."),
    ],
    [2300, 7060],
    font_size=8.2,
)
add_heading("13.2 Common Report Filters", 2)
add_p(
    "Workspace, reporting period, original/base currency, owner, vendor, tool, category/subcategory, subscription type, status, "
    "billing cadence, payment account, cost entry type/status, target scope/severity, tag, renewal window, amount range, data source, "
    "reconciliation state, and freshness."
)
add_heading("13.3 Notification Catalogue", 2)
add_table(
    ["Event", "Default timing / trigger", "Minimum message content"],
    [
        ("Target warning", "Warning basis reached", "Scope, period, current/forecast, target, variance, contributors, action link."),
        ("Target critical", "Critical basis reached or warning escalates", "Severity, current/forecast, target, gap, owner/assignee, action link."),
        ("Renewal due", "Configurable 30/14/7/1 days", "Tool, plan, amount/currency, renewal, cancellation deadline, payment alias."),
        ("Trial ending", "Configurable 7/3/1 days", "Trial end, expected conversion price, auto-convert status, action choices."),
        ("Cancellation deadline", "Before contractual deadline", "Deadline, renewal amount, decision status and direct review link."),
        ("Payment account expiry", "90/60/30 days or conflict detected", "Masked account, expiry, affected subscriptions and next charges."),
        ("Token/credit low", "Balance or days-of-use threshold", "Remaining units, burn rate, expiry, forecast, package/action link."),
        ("Unusual usage", "Rate-of-change/anomaly rule", "Period, baseline, current usage/cost, source freshness and contributors."),
        ("Stale data / connector failure", "Expected refresh missed or repeated failure", "Source, last success, missing period, retry/manual action."),
        ("Import/export/report complete", "Job terminal state", "Outcome counts, safe file link/expiry or error summary."),
    ],
    [1800, 2700, 4860],
    font_size=7.8,
)
add_heading("13.4 Alert Severity and State", 2)
add_table(
    ["Dimension", "Values", "Rules"],
    [
        ("Severity", "Info, Warning, Critical", "Severity escalates when risk increases; recovery does not erase prior event."),
        ("Lifecycle", "Open, Acknowledged, Snoozed, Assigned, Resolved, Dismissed, Auto-recovered", "Only permitted transitions; dismiss requires reason for critical alerts."),
        ("Delivery", "Pending, Sent, Delivered, Failed, Retrying, Suppressed", "Per-channel status retained; quiet hours delay, not lose, eligible messages."),
        ("Deduplication", "Open event key + cool-down", "Same scope/type/period is grouped; higher severity may notify immediately."),
    ],
    [1500, 2500, 5360],
    font_size=8.2,
)

# 7 Detailed use cases
add_heading("7. Detailed Use Cases", 1)
add_p(
    "Use cases describe user intent and externally observable behaviour. Screen names are indicative; the final information architecture "
    "may consolidate steps while preserving all validations and outcomes."
)

add_use_case(
    "UC-01", "Complete first-time setup", "Account Owner",
    "Create a usable workspace and cost-control baseline.",
    "User has verified an account and has not completed onboarding.",
    "User signs in for the first time.",
    [
        "System asks for base currency, time zone, locale, fiscal-month start, and notification preferences.",
        "User optionally loads safe sample data or chooses a blank workspace.",
        "User adds at least one payment account alias without entering sensitive card credentials.",
        "User adds or imports the first subscription and selects its type.",
        "System previews normalized cost, next renewal, and recommended starter target.",
        "User confirms settings and reaches the dashboard with an onboarding checklist.",
    ],
    [
        "User skips optional steps; system retains checklist and uses disclosed defaults.",
        "Import contains invalid rows; system shows row-level errors without creating partial ambiguous records.",
    ],
    "Preferences and valid records are saved; onboarding state is recorded.",
    "FR-003, FR-020, FR-051, FR-060, FR-083, FR-084",
)

add_use_case(
    "UC-02", "Register a monthly or annual subscription", "Owner / Editor",
    "Track a recurring tool and its future obligations.",
    "Tool and subscription type exist; user has edit permission.",
    "User selects Add subscription.",
    [
        "User selects or creates vendor and tool.",
        "User selects monthly, annual, or custom recurring type.",
        "User enters plan, amount, currency, billing anchor, renewal, auto-renew, cancellation deadline, and tax/discount.",
        "User selects payment account or records a permitted payment exception.",
        "System shows cash schedule, normalized monthly/annual cost, and reminder dates.",
        "User confirms; system creates subscription, plan version, commitments, audit entry, and alert events.",
    ],
    [
        "Billing date is invalid for a month; system applies the disclosed end-of-month rule.",
        "A likely duplicate exists; system warns and offers view/continue.",
    ],
    "Active subscription is visible in lists, forecast, calendar, payment account, and reports.",
    "FR-010–FR-014, FR-020–FR-023, FR-028, FR-032, FR-053",
)

add_use_case(
    "UC-03", "Register an on-demand AI service", "Owner / Editor",
    "Track variable usage and prevent budget overrun.",
    "Tool exists and on-demand or hybrid type is active.",
    "User selects an on-demand subscription model.",
    [
        "User selects meter units and defines flat/tiered/package or hybrid rates.",
        "User enters included units, minimums, caps, billing period, and data source.",
        "User selects forecast method and target amount.",
        "User configures warning, critical, rapid-growth, low-balance, and stale-data alerts.",
        "System validates rate continuity, displays examples, and saves the configuration.",
    ],
    [
        "Rate tiers have gaps/overlap; system blocks activation and highlights the affected range.",
        "No target is entered; system allows activation only after the user explicitly accepts no budget alert.",
    ],
    "Usage can be posted and cost/forecast evaluated using the effective pricing configuration.",
    "FR-040–FR-047, FR-060–FR-067",
)

add_use_case(
    "UC-04", "Record token package purchase and consumption", "Owner / Editor",
    "Maintain an auditable balance and effective unit cost.",
    "Prepaid token/credit subscription exists.",
    "User records a package purchase or imports usage.",
    [
        "User enters package units, paid/free split, net cost, currency, purchase date, and expiry.",
        "System posts the purchase and creates balance lots.",
        "User enters or imports consumption by service/model and period.",
        "System detects duplicates, applies included units and FIFO-by-expiry allocation, and calculates cost.",
        "User reviews allocations, remaining balance, effective unit cost, and expiry risk.",
    ],
    [
        "Consumption exceeds available balance; system records eligible overage or holds the excess as an exception.",
        "User overrides allocation; system requires a reason and writes audit history.",
    ],
    "Balance ledger and cost ledger reconcile; alerts are evaluated.",
    "FR-031, FR-040–FR-047, BR-007, BR-017",
)

add_use_case(
    "UC-05", "Change plan or price", "Owner / Editor",
    "Apply a future or immediate plan change without corrupting history.",
    "Subscription is active/trial/paused and user can edit.",
    "Vendor changes price, cadence, currency, entitlements, or payment account.",
    [
        "User opens Change plan and enters effective date and revised terms.",
        "System validates date overlap and previews affected commitments, targets, and forecast.",
        "User confirms; system closes the previous version and opens the new version.",
        "Future expected commitments and alerts are regenerated.",
        "System records before/after summary and reason in the audit trail.",
    ],
    [
        "Change is backdated into a closed period; system blocks or routes to authorized correction workflow.",
        "A posted charge differs from expected price; system preserves it and flags reconciliation variance.",
    ],
    "Historical records use old version; current/future calculations use the new version.",
    "FR-023, FR-028, FR-032, FR-036, FR-063, FR-081",
)

add_use_case(
    "UC-06", "Record and reconcile a charge", "Owner / Finance Reviewer",
    "Post an actual cost and match it to the correct obligation.",
    "Subscription or unmatched-cost queue exists.",
    "User enters/imports a charge.",
    [
        "User provides transaction/service dates, amount, tax, discount, currency, payment account, and reference.",
        "System proposes subscription and expected commitment matches.",
        "User accepts or changes match and reviews FX conversion.",
        "System posts the entry, offsets the matched commitment, and recalculates totals.",
        "If target severity changes, the system creates or escalates an alert.",
    ],
    [
        "Duplicate reference/amount/date detected; system blocks or requires explicit override.",
        "No match exists; entry remains posted/unmatched and appears in reconciliation exceptions.",
    ],
    "Ledger, commitment, payment account, forecast, target, and audit views are updated once.",
    "FR-030–FR-036, FR-060–FR-066",
)

add_use_case(
    "UC-07", "Create a target and alert policy", "Owner / Editor",
    "Control spending for a defined scope and period.",
    "At least one eligible scope exists.",
    "User selects Set target.",
    [
        "User selects global/category/tool/subscription/token/payment-account scope.",
        "User selects period, amount, currency, comparison basis, and carry-forward rule.",
        "User enters warning and critical thresholds and chooses channels/recipients.",
        "System previews current utilization and example alert timing.",
        "User activates the target and policy.",
    ],
    [
        "Thresholds are reversed or outside allowed range; system blocks save.",
        "An overlapping target exists; system explains overlap and allows intentional coexistence.",
    ],
    "Target appears on dashboards and evaluates whenever relevant cost data changes.",
    "FR-060–FR-067",
)

add_use_case(
    "UC-08", "Respond to a target-crossing alert", "Owner / Assignee",
    "Understand and resolve a spend risk.",
    "Active target crossed a configured threshold.",
    "Alert is delivered in-app or by email.",
    [
        "User opens the alert and sees severity, amount, variance, forecast, scope, and contributing costs.",
        "User drills into underlying subscriptions/usage and calculation explanation.",
        "User chooses acknowledge, assign, snooze, edit target, correct data, reduce usage, or cancel/change a plan.",
        "System records the action and reevaluates severity after any data change.",
        "User resolves the alert with a note when the risk is addressed or formally accepted.",
    ],
    [
        "Email delivery fails; in-app alert remains active and retry status is visible.",
        "Forecast drops below threshold; system records automatic recovery but retains the event history.",
    ],
    "Alert state, action history, and target status are updated without duplicate notifications.",
    "FR-063–FR-066, FR-074, FR-081",
)

add_use_case(
    "UC-09", "Review renewal and cancel a subscription", "Owner / Editor",
    "Avoid an unwanted renewal and retain evidence.",
    "Active subscription has renewal/cancellation dates.",
    "Renewal reminder is received or calendar item opened.",
    [
        "User reviews price, utilization, recent cost, forecast, payment account, and cancellation deadline.",
        "User records decision to cancel, effective/access-end date, reason, and optional evidence attachment.",
        "System previews avoided future cost, outstanding fees, and affected alerts.",
        "User confirms cancellation.",
        "System updates status/schedule, stops eligible future commitments, retains history, and records outcome.",
    ],
    [
        "Cancellation is after deadline; system warns that renewal may still occur and keeps the commitment until confirmed.",
        "Vendor issues refund/fee; user posts linked ledger entry.",
    ],
    "Subscription lifecycle and forecast reflect the confirmed commercial outcome.",
    "FR-021, FR-025, FR-027, FR-028, FR-033, FR-064–FR-066",
)

add_use_case(
    "UC-10", "Replace an expiring payment account", "Owner / Administrator",
    "Move future charges safely while retaining payment history.",
    "Existing account is expiring/inactive; replacement account exists.",
    "Expiry alert or account action is opened.",
    [
        "System lists affected subscriptions, next charges, currencies, and renewal dates.",
        "User selects all or a subset and chooses replacement effective date.",
        "System validates replacement status/currency and previews impact.",
        "User confirms; future plan versions reference the replacement account.",
        "Original account is marked inactive when no future obligations remain.",
    ],
    [
        "Some subscriptions must be updated at vendor portal; system creates follow-up actions rather than claiming completion.",
        "User lacks access to a linked subscription; bulk update excludes it and reports the exception.",
    ],
    "Future charges reference the new account; historical charges retain the old masked account.",
    "FR-050–FR-055, FR-081",
)

add_use_case(
    "UC-11", "Import subscriptions, usage or costs", "Owner / Editor",
    "Create or update records efficiently from a file.",
    "User has a supported file and import permission.",
    "User opens Import.",
    [
        "User selects entity type, downloads template or uploads CSV/XLSX, and chooses a saved mapping.",
        "System scans structure, shows mapping, normalizes values, and previews valid/warning/error rows.",
        "User resolves required mappings and duplicate decisions.",
        "System executes the approved batch and reports created, updated, skipped, and failed rows.",
        "User downloads error file and can rerun corrected rows using the batch reference.",
    ],
    [
        "Fatal validation or malware scan failure prevents processing.",
        "Background job fails; system leaves batch retryable and avoids duplicate committed rows.",
    ],
    "Import batch, row outcomes, audit events, and recalculated totals are available.",
    "FR-020, FR-034, FR-042, FR-084, FR-085",
)

add_use_case(
    "UC-12", "Analyze dashboard and drill into cost", "Owner / Viewer",
    "Understand current cost position and the reason for each total.",
    "User has view permission and records exist.",
    "User opens Dashboard or a saved view.",
    [
        "System loads actual, committed, forecast, targets, alerts, renewals, token usage, and payment risks for the selected period.",
        "User changes period, currency view, owner/category/status filters, or saved view.",
        "Charts and KPIs update consistently and preserve filter context.",
        "User selects a KPI/chart segment to open the contributing records and calculation explanation.",
        "User exports or shares a masked report if authorized.",
    ],
    [
        "No data exists; system shows a helpful empty state and onboarding actions.",
        "Some source data is stale; system marks freshness and excludes/estimates according to disclosed rules.",
    ],
    "User can explain the total and take a relevant follow-up action.",
    "FR-070–FR-077",
)

add_use_case(
    "UC-13", "Correct a posted historical cost", "Finance Reviewer / Administrator",
    "Correct an error while preserving financial and audit integrity.",
    "Posted entry exists and user has correction permission.",
    "User selects Correct entry.",
    [
        "System displays original entry and dependent matches, totals, and targets.",
        "User enters correction reason and replacement values.",
        "System creates a linked reversing entry and a new corrected entry.",
        "System recalculates affected periods, snapshots, targets, and alerts according to close rules.",
        "Audit view shows the entire correction chain.",
    ],
    [
        "Period is closed; authorized user must reopen with reason or post an adjustment in the current period.",
        "Correction would unmatch a refund/credit; system requires resolving the dependency first.",
    ],
    "No original evidence is lost; reports use the valid posted chain.",
    "FR-036, FR-037, FR-081",
)

add_use_case(
    "UC-14", "Manage subscription types and reference data", "Administrator",
    "Keep configuration relevant without breaking existing data.",
    "Administrator is authenticated.",
    "Admin opens Settings / Reference data.",
    [
        "Administrator creates or edits a subscription type/category with display order, active state, and applicable cost characteristics.",
        "System validates unique keys and configuration consistency.",
        "Administrator previews affected records before archive or merge.",
        "System blocks destructive deletion and offers replacement mapping.",
        "System logs the change and applies it to new/eligible forms.",
    ],
    [
        "A type is locked because it is system-required; admin may rename display label but not alter core semantics.",
        "Mapping would create invalid data; system reports affected records and stops.",
    ],
    "Reference configuration is active, traceable, and backward compatible.",
    "FR-012–FR-015, FR-080–FR-082",
)

add_use_case(
    "UC-15", "Export data and close the account", "Account Owner",
    "Obtain a portable copy and deactivate service safely.",
    "Owner is authenticated and passes step-up verification if configured.",
    "User opens Privacy / Account.",
    [
        "User requests a data export and selects included scopes.",
        "System prepares a masked, time-limited export and records the event.",
        "User reviews active renewals, retention, legal hold, and consequences.",
        "User confirms deactivation/closure with reason and verification.",
        "System revokes active sessions, stops notifications/jobs as appropriate, and applies retention workflow.",
    ],
    [
        "Export job fails; user is notified and may retry without duplicate files.",
        "Retention hold prevents purge; system explains the retained categories and deactivation continues.",
    ],
    "Account is deactivated, data disposition is scheduled, and evidence is retained as authorized.",
    "FR-006, FR-075, FR-081, FR-085, FR-087",
)

page_break()

# 1-3 Business framing
add_heading("1. Executive Summary", 1)
add_p(
    "Individuals and small teams increasingly pay for a mix of AI chat products, model APIs, token packages, "
    "developer utilities, design products, storage, automation platforms, and other digital tools. Charges may be "
    "monthly, yearly, prepaid, usage-based, trial-based, or hybrid. They may be billed in different currencies and "
    "to different cards or accounts. Without a single control point, renewals are missed, unused products remain "
    "active, token spending exceeds expectations, and the true monthly cost is difficult to understand."
)
add_p(
    "The proposed Personal AI Subscription Management Platform will provide one source of truth for subscriptions, "
    "usage and payment responsibility. It will normalize recurring and on-demand costs into monthly and annual views, "
    "track planned targets against actual and forecast spend, warn users before thresholds or renewals are crossed, "
    "and preserve a complete cost and lifecycle history."
)
add_heading("1.1 Proposed Business Capability", 2)
add_bullets([
    "Register and manage AI subscriptions, AI token/API plans, other software tools, and custom subscription types.",
    "Support free, trial, monthly, quarterly, annual, prepaid-credit, on-demand, and hybrid charging models.",
    "Track recurring fees, add-ons, taxes, discounts, token purchases, metered usage, credits, refunds, and manual adjustments.",
    "Manage masked payment accounts such as debit/credit cards, bank accounts, wallets, virtual cards, and vendor balances.",
    "Set global, category, tool, and on-demand targets with warning and critical thresholds.",
    "Forecast end-of-period spend and projected renewal obligations using configurable costing rules.",
    "Send in-app and email alerts for target crossings, renewal windows, payment expiry, trial expiry, unusual usage, and missing data.",
    "Provide compact dashboards, responsive workflows, accessible forms, reports, exports, and an auditable history.",
])
add_heading("1.2 Business Value", 2)
add_table(
    ["Outcome", "Expected value"],
    [
        ("Visibility", "A reliable monthly, annual, category, tool, and payment-source view of commitments and usage."),
        ("Cost control", "Early warning before spend exceeds target and clear identification of unused or duplicative tools."),
        ("Renewal control", "Reduced accidental renewal, expired trials, service interruption, and late cancellation."),
        ("Governance", "Traceable ownership, payment responsibility, approval status, changes, and exported evidence."),
        ("Decision support", "Comparable effective monthly costs, cost trends, utilization, and savings opportunities."),
    ],
    [2000, 7360],
    font_size=8.8,
)

add_heading("2. Business Context, Objectives and Success Measures", 1)
add_heading("2.1 Problem Statement", 2)
add_p(
    "Subscription records are commonly dispersed across inboxes, vendor portals, bank statements, spreadsheets, and memory. "
    "A charge amount alone does not show whether it is a monthly commitment, annual renewal, token top-up, or on-demand usage. "
    "This fragmentation makes budgeting reactive and creates avoidable financial and operational risk."
)
add_heading("2.2 Objectives", 2)
add_table(
    ["ID", "Objective", "Business measure"],
    [
        ("OBJ-01", "Create one complete inventory of active, trial, paused, cancelled, and expired tools.", "At least 95% of known subscriptions recorded after onboarding."),
        ("OBJ-02", "Show actual, committed, and forecast costs by period.", "Monthly totals reconcile to source data within approved tolerance."),
        ("OBJ-03", "Prevent surprise overages and renewals.", "Threshold and renewal notifications sent within configured windows."),
        ("OBJ-04", "Make cost ownership and payment source clear.", "Every paid subscription linked to an owner and active payment account or documented exception."),
        ("OBJ-05", "Support cost optimization decisions.", "Dashboard identifies unused, low-use, duplicate, or high-growth subscriptions."),
        ("OBJ-06", "Provide a fast, usable experience on desktop and mobile.", "Core tasks complete without horizontal scrolling at supported mobile widths."),
    ],
    [850, 4200, 4310],
    font_size=8.2,
)
add_heading("2.3 Key Performance Indicators", 2)
add_bullets([
    "Inventory coverage: registered paid and trial tools divided by known tools.",
    "Spend variance: actual month-to-date and forecast end-of-month versus target.",
    "Prevented renewal value: value of subscriptions cancelled before unwanted renewal.",
    "Reconciliation rate: cost entries matched to a subscription and payment account.",
    "Notification effectiveness: alerts delivered, acknowledged, snoozed, and acted upon.",
    "Data freshness: percentage of active usage-based subscriptions updated within their expected cycle.",
    "Task efficiency: median time to add a subscription, record a usage charge, and identify the next renewal.",
])
add_heading("2.4 Definition of Success", 2)
add_callout(
    "MINIMUM SUCCESS",
    "A user can register all subscription models, see normalized monthly cost and forecast, set targets, receive actionable alerts, "
    "identify the card/account used, and produce an accurate cost report from any supported device.",
    fill=PALE_GREEN,
    color=GREEN,
)

add_heading("3. Scope, Stakeholders, Assumptions and Constraints", 1)
add_heading("3.1 In Scope", 2)
add_bullets([
    "User authentication, profile, currency, locale, time-zone, notification, and dashboard preferences.",
    "Tool/vendor catalogue and user-defined categories, tags, ownership, purpose, environment, and status.",
    "Subscription type configuration and lifecycle management for recurring, prepaid, usage-based, and hybrid plans.",
    "AI token/API pricing, packages, credits, usage entries, budget targets, forecast, and effective-unit cost.",
    "Payment-account register with masked details, expiry, status, billing currency, owner, and linked subscriptions.",
    "Cost ledger, taxes, discounts, refunds, credits, adjustments, proration, exchange-rate capture, and attachments.",
    "Targets, threshold policies, alert delivery, renewal reminders, trial reminders, unusual-cost detection, and acknowledgement.",
    "Dashboards, calendars, searchable lists, comparison, cost trends, reports, CSV/XLSX export, and printable summaries.",
    "Administrative reference data, role permissions, audit history, background jobs, system health, backup, and retention.",
    "Responsive web interface built with Laravel, MySQL, and a licensed Metronic template package.",
])
add_heading("3.2 Out of Scope for the Initial Release", 2)
add_bullets([
    "Initiating or settling vendor payments, changing a vendor subscription directly, or storing full card numbers/CVV.",
    "Automatic scraping of arbitrary vendor portals without an approved API or user-provided import.",
    "Corporate procurement, purchase orders, complex multi-entity accounting, tax filing, or general-ledger posting.",
    "Native iOS/Android applications; the first release is a responsive web application and may later become a PWA.",
    "Real-time bank aggregation unless delivered as a separately approved integration.",
    "Vendor-specific token measurement where the vendor does not expose usage data; manual entry/import remains supported.",
])
add_heading("3.3 Stakeholders and Roles", 2)
add_table(
    ["Stakeholder / role", "Primary interest", "Responsibilities"],
    [
        ("Business Sponsor", "Value, funding, delivery outcome", "Approve scope, budget, priorities, and go-live."),
        ("Product Owner", "Requirement quality and value", "Own backlog, rules, acceptance, and stakeholder decisions."),
        ("Account Owner / User", "Personal subscription control", "Maintain tools, targets, payment sources, and actions."),
        ("Viewer / Auditor", "Read-only evidence", "Review dashboard, reports, costs, and history."),
        ("Administrator", "Configuration and support", "Manage reference data, users, roles, jobs, templates, and support actions."),
        ("Finance Reviewer (optional)", "Cost accuracy", "Review reconciliation, adjustments, reports, and exceptions."),
        ("Engineering Team", "Feasible implementation", "Design, build, migrate, operate, and document the platform."),
        ("QA / Security", "Quality and control", "Verify requirements, security, accessibility, performance, and recovery."),
    ],
    [2100, 2700, 4560],
    font_size=8.2,
)
add_heading("3.4 Assumptions", 2)
add_bullets([
    "The first deployment is single-tenant or personal-workspace oriented; data structures should permit future multi-user workspaces.",
    "The user chooses one base reporting currency; transaction and subscription currencies are retained separately.",
    "Exchange rates may be entered manually or obtained from an approved provider; the rate, source, and effective date are stored.",
    "Notification delivery initially includes in-app and email; optional browser push, SMS, or chat channels require separate credentials.",
    "A subscription may have multiple cost components and plan changes over time; financial history is immutable except by controlled adjustment.",
    "The system is a management and forecasting tool, not a legal, accounting, or payment-processing system.",
])
add_heading("3.5 Constraints and Dependencies", 2)
add_bullets([
    "Development stack: supported PHP and Laravel release selected at delivery start, MySQL 8-compatible database, Metronic assets, queue workers, scheduler, and email provider.",
    "Use of Metronic is subject to the purchaser holding the correct commercial license; no third-party credentials or paid assets are embedded in source control.",
    "Vendor APIs differ in availability and billing semantics; every automated connector must degrade safely to manual entry or file import.",
    "Payment account data must be masked and minimized. Full PAN, CVV, online-banking password, private API key, and secret token storage are prohibited.",
    "Final tax, privacy, retention, and accessibility obligations depend on the deployment jurisdiction and must be confirmed before production.",
])

page_break()

# 4 Business concepts/rules
add_heading("4. Business Concepts, Rules and Cost Calculations", 1)
add_heading("4.1 Core Definitions", 2)
add_table(
    ["Term", "Definition"],
    [
        ("Tool", "An AI or non-AI software product whose access, use, or cost is managed."),
        ("Vendor", "The party billing for a tool; one vendor may supply multiple tools."),
        ("Subscription", "A time-bounded commercial arrangement for a tool, including free, trial, recurring, prepaid, on-demand, or hybrid models."),
        ("Plan / rate version", "The price and entitlement configuration effective for a defined date range."),
        ("Cost entry", "A ledger event such as recurring fee, usage charge, token purchase, tax, discount, credit, refund, or adjustment."),
        ("Token package", "A purchased or granted balance of units/credits/tokens with optional expiry."),
        ("Payment account", "A masked card, bank account, wallet, virtual card, or vendor balance used to pay a charge."),
        ("Target", "An approved spending ceiling or reference amount for a period, scope, and currency."),
        ("Alert policy", "Thresholds, time windows, channels, recipients, quiet hours, and repetition rules."),
        ("Committed cost", "A known contractual cost expected in the selected period regardless of usage."),
        ("Forecast cost", "Projected total cost for the period using actuals plus future commitments and estimated on-demand usage."),
    ],
    [1800, 7560],
    font_size=8.4,
)
add_heading("4.2 Subscription Type Model", 2)
add_table(
    ["Type", "Required characteristics", "Cost treatment"],
    [
        ("Free", "No charge; optional feature/usage tracking", "Zero committed cost."),
        ("Trial", "Start, end, post-trial action, expected conversion price", "Forecast may include post-trial fee if conversion is enabled."),
        ("Monthly recurring", "Billing anchor, amount, renewal, proration rule", "Full or prorated fee in each billing period."),
        ("Annual recurring", "Renewal date, annual fee, optional installment plan", "Cash view at renewal; normalized monthly view equals annual net cost / 12."),
        ("Other recurring", "Interval unit and count, e.g., quarterly or every 4 weeks", "Occurrences generated by recurrence schedule."),
        ("Prepaid token / credit", "Package units, purchase cost, expiry, balance", "Cash at purchase; consumption and effective-unit cost tracked separately."),
        ("On-demand", "Meter, rate tiers, minimums, caps, usage source", "Actual usage x applicable rate plus fixed/minimum charges."),
        ("Hybrid", "Recurring base fee plus included units and overage tiers", "Base commitment plus forecast/actual overage."),
        ("One-time", "Purchase date and amount", "Recognized once; optional amortized management view."),
    ],
    [1500, 4300, 3560],
    font_size=8.1,
)
add_heading("4.3 Business Rules", 2)
business_rules = [
    ("BR-001", "Every subscription belongs to exactly one tool and has one active type at a time."),
    ("BR-002", "Plan, price, currency, billing cadence, tax treatment, and included entitlements are versioned by effective date."),
    ("BR-003", "A paid subscription must reference a payment account unless it is marked reimbursed, invoice/manual, vendor credit, or exception with reason."),
    ("BR-004", "Payment account identifiers display only friendly name, issuer/provider, type, last four digits or alias, owner, expiry month/year, and status."),
    ("BR-005", "Historical cost entries are not silently recalculated when a current price, target, or exchange rate changes."),
    ("BR-006", "The system stores original amount/currency and base-currency amount/rate/source for each financial event."),
    ("BR-007", "Refunds, credits, and reversals reference the original entry where possible and never delete the audit record."),
    ("BR-008", "A target has a scope, period, amount, currency, warning threshold, critical threshold, and status."),
    ("BR-009", "Threshold comparisons use the selected basis: actual, committed, forecast, or actual-plus-committed."),
    ("BR-010", "A threshold event is deduplicated for the configured cool-down window but may escalate when a higher threshold is crossed."),
    ("BR-011", "On-demand forecast uses a user-selected method: fixed target, run-rate, recent-period average, manual forecast, or zero until usage is recorded."),
    ("BR-012", "Trial and renewal dates are evaluated in the workspace time zone, with reminders scheduled by calendar date rather than server time."),
    ("BR-013", "Pausing stops future recurring forecast only when the vendor contract is also paused; local status alone does not erase commitments."),
    ("BR-014", "Cancellation retains historical costs and may include an access-end date, cancellation deadline, refund, or cancellation fee."),
    ("BR-015", "Deleted records with financial or audit history are archived/soft-deleted; only authorized administrators may purge after retention checks."),
    ("BR-016", "Cost totals exclude drafts and voided entries, include posted adjustments, and expose tax-inclusive and tax-exclusive views."),
    ("BR-017", "When included token units reset, unused carryover is applied only if the plan version explicitly allows it."),
    ("BR-018", "Shared workspace totals must respect row-level ownership/visibility and never leak one user's masked payment data to another."),
    ("BR-019", "All dates, currencies, and numeric inputs are validated server-side; browser formatting is not authoritative."),
    ("BR-020", "Alert delivery failure is recorded and retried using controlled backoff; a failed delivery does not mark an alert acknowledged."),
]
add_table(
    ["Rule ID", "Rule"],
    business_rules,
    [1000, 8360],
    font_size=8.3,
)
add_heading("4.4 Cost Calculation Rules", 2)
add_table(
    ["Metric", "Calculation / interpretation"],
    [
        ("Normalized monthly recurring cost", "Net recurring amount × 12 / billing interval in months; annual plans divide by 12."),
        ("Normalized annual recurring cost", "Normalized monthly recurring cost × 12."),
        ("Net cost entry", "Gross charge + tax + surcharge − discount − credit − refund."),
        ("On-demand actual", "Posted metered usage evaluated against the effective pricing tier plus fixed/minimum charges."),
        ("Run-rate forecast", "Actual period-to-date / elapsed billable days × total billable days, bounded by configured minimum/maximum if present."),
        ("Hybrid forecast", "Base fee + forecast overage after included units and carryover."),
        ("Forecast period total", "Actual posted + expected committed not yet posted + forecast variable usage + expected taxes/fees."),
        ("Target utilization", "Comparison basis amount / target amount × 100; undefined when target equals zero."),
        ("Effective token cost", "Allocated net package cost / consumed paid units; free/granted units shown separately."),
        ("Savings opportunity", "Projected avoidable future cost after cancellation deadline, excluding sunk cost and mandatory fees."),
        ("FX conversion", "Original amount × stored conversion rate to base currency; inverse rate is clearly identified if used."),
    ],
    [2300, 7060],
    font_size=8.25,
)
add_callout(
    "ROUNDING",
    "Perform calculations using database decimals at provider-supported precision. Round only for display or final currency posting. "
    "Totals must equal the sum of stored base-currency entries within a configurable minor-unit tolerance.",
    fill=PALE_AMBER,
    color=AMBER,
)

add_heading("5. Future-State Processes", 1)
add_heading("5.1 Subscription Lifecycle", 2)
add_p(
    "Discover / import → validate vendor and tool → select subscription type → record plan and payment account → "
    "set targets and alert policy → activate / trial → record charges and usage → review forecast → renew, change, pause, or cancel → retain history."
)
add_heading("5.2 Monthly Cost-Control Cycle", 2)
add_numbers([
    "At period start, create expected commitments from active plan schedules and carry forward approved targets.",
    "Ingest or manually record recurring fees, token purchases, usage, credits, refunds, taxes, and adjustments.",
    "Reconcile entries to subscriptions and payment accounts; route unmatched entries to an exception queue.",
    "Recalculate actual, committed, and forecast totals by category, tool, payment source, and target scope.",
    "Evaluate warning and critical thresholds; create deduplicated alerts and escalation tasks.",
    "Review the dashboard, acknowledge alerts, correct missing data, and take renewal/cancellation actions.",
    "Close the period, lock the management snapshot, and compare actual versus target and prior periods.",
])
add_heading("5.3 Renewal and Trial Review", 2)
add_numbers([
    "Scheduler identifies active trials, renewal/cancellation deadlines, price reviews, and payment expiries in configured windows.",
    "System issues reminders with tool, amount, payment account alias, deadline, forecast impact, and direct action link.",
    "User chooses renew, change plan, cancel, pause, snooze, or record a decision note.",
    "The plan schedule and forecast update from the effective date; historical entries remain unchanged.",
])
add_heading("5.4 On-Demand and Token Control", 2)
add_numbers([
    "Define meter unit, pricing tiers, included units, budget target, forecast method, and data-refresh expectation.",
    "Record usage manually, via file import, or through an approved connector; detect duplicates by provider reference and period.",
    "Apply included units, carryover, package balances, tier rates, minimums, and caps in effective-date order.",
    "Compare consumption and cost against target; alert on rapid growth, stale data, depleted balance, or forecast overrun.",
])

page_break()

# 6 Functional requirements
add_heading("6. Functional Requirements", 1)
add_p(
    "Priority uses MoSCoW: Must is required for the initial production release; Should is strongly expected; Could may be scheduled after core acceptance. "
    "Each requirement is testable and traceable to at least one use case or acceptance scenario."
)

add_requirement_group("6.1 Identity, Workspace and Preferences", [
    ("FR-001", "Authenticate with email/password and optional approved social/SSO provider.", "Must", "Valid users sign in; invalid and locked users receive safe errors."),
    ("FR-002", "Support email verification, password reset, session logout, and configurable inactivity timeout.", "Must", "Security flows complete with expiring, single-use tokens."),
    ("FR-003", "Maintain profile, base currency, locale, time zone, date/number format, fiscal month, and week start.", "Must", "All views and scheduled actions honor saved preferences."),
    ("FR-004", "Support owner, administrator, editor, viewer/auditor, and optional finance-reviewer roles.", "Must", "Each protected action is permitted or denied server-side by role."),
    ("FR-005", "Allow dashboard, default period, column, density, and notification preferences.", "Should", "Preferences persist across devices and can be reset."),
    ("FR-006", "Provide account data export and controlled deactivation/closure workflow.", "Should", "User receives export and retention warning before closure."),
])

add_requirement_group("6.2 Tool, Vendor, Category and Subscription Type Management", [
    ("FR-010", "Create, edit, archive, search, filter, tag, and merge duplicate vendors and tools.", "Must", "Records preserve linked subscription and history references."),
    ("FR-011", "Store tool name, vendor, category, purpose, URL, logo/icon, notes, owner, environment, tags, and status.", "Must", "Required fields validate and display consistently."),
    ("FR-012", "Manage configurable categories and subcategories with color/icon and active state.", "Must", "Categories can be reused in filters, targets, and reports."),
    ("FR-013", "Manage subscription types and characteristics: recurring, prepaid, usage, trial, free, one-time, hybrid.", "Must", "Type rules drive relevant form fields and calculations."),
    ("FR-014", "Prevent deletion of a type/category in use; allow replacement mapping and archive.", "Must", "System offers remap or blocks with impact count."),
    ("FR-015", "Support custom fields for tool/subscription metadata with type, validation, visibility, and order.", "Could", "Configured fields render and validate without code change."),
])

add_requirement_group("6.3 Subscription Lifecycle and Plan Versioning", [
    ("FR-020", "Create subscriptions manually, by duplication, guided onboarding, or validated import.", "Must", "A complete active subscription can be saved in under three screens."),
    ("FR-021", "Capture plan name, status, start/end, trial end, billing cadence, billing anchor, renewal, cancellation deadline, grace period, and auto-renew.", "Must", "Lifecycle dates produce correct reminders and schedule."),
    ("FR-022", "Support monthly, annual, interval-based, prepaid, on-demand, hybrid, free, trial, and one-time pricing.", "Must", "Relevant calculations match the selected model."),
    ("FR-023", "Version plan price, currency, entitlements, included units, tax, discount, and payment account by effective date.", "Must", "Past costs retain prior version; future forecast uses current version."),
    ("FR-024", "Support add-ons, seats, environments, and multiple cost components under one subscription.", "Should", "Components roll into totals and remain individually reportable."),
    ("FR-025", "Support pause, resume, cancel, expire, reactivate, and archive with reason and effective date.", "Must", "Status transition rules prevent invalid or backdated overlap."),
    ("FR-026", "Record vendor account reference, login URL, owner, recovery contact, and attachment without storing passwords.", "Should", "Sensitive secret fields are not available."),
    ("FR-027", "Maintain notes, decision history, cancellation evidence, invoice/receipt attachments, and renewal outcome.", "Should", "Authorized users can retrieve the evidence from history."),
    ("FR-028", "Show next charge, renewal, cancellation deadline, days remaining, and normalized cost.", "Must", "Values update after any schedule or plan change."),
])

add_requirement_group("6.4 Cost Ledger, Recurring Charges and Adjustments", [
    ("FR-030", "Maintain a ledger of draft, posted, voided, refunded, and reconciled cost entries.", "Must", "Only posted/eligible entries affect totals per business rules."),
    ("FR-031", "Record charge type, service period, transaction date, due date, amount, tax, discount, credit, currency, FX rate/source, vendor reference, and notes.", "Must", "Entry preserves original and base-currency values."),
    ("FR-032", "Generate expected recurring commitments from subscription schedules without duplicating posted charges.", "Must", "Matched actuals replace/offset expected commitments once."),
    ("FR-033", "Support proration, partial periods, minimum charges, setup fees, late fees, refunds, credits, and manual adjustments.", "Must", "Net totals and audit references calculate correctly."),
    ("FR-034", "Import cost entries from CSV/XLSX with mapping, preview, validation, duplicate detection, and error export.", "Should", "Valid rows import atomically or by clear partial-result policy."),
    ("FR-035", "Reconcile an entry to tool, subscription, cost component, payment account, and statement reference.", "Should", "Matched and unmatched totals are visible."),
    ("FR-036", "Permit authorized correction by reversal and replacement, not silent historical overwrite.", "Must", "Original, reversal, replacement, actor, and reason are linked."),
    ("FR-037", "Close and reopen reporting periods under permission with a reason.", "Should", "Closed periods block ordinary edits and preserve snapshot totals."),
])

add_requirement_group("6.5 AI Token, Credit and On-Demand Usage", [
    ("FR-040", "Define meter units such as input tokens, output tokens, requests, images, minutes, credits, or custom units.", "Must", "Units retain precision and are labeled in all views."),
    ("FR-041", "Define flat, tiered, volume, package, included-unit, minimum, capped, and hybrid rates by effective date.", "Must", "Pricing engine applies the correct rate version and tier logic."),
    ("FR-042", "Record usage by model/service, environment/project, date range, unit type, quantity, and provider reference.", "Must", "Duplicates are detected and totals can be segmented."),
    ("FR-043", "Track token/credit packages, granted units, consumed units, remaining balance, carryover, and expiry.", "Must", "Balance ledger reconciles purchases, grants, use, and expiry."),
    ("FR-044", "Allocate package consumption using configurable FIFO-by-expiry default and manual override with reason.", "Should", "Allocation and effective cost are reproducible."),
    ("FR-045", "Calculate actual usage cost, effective unit cost, included-unit utilization, and estimated overage.", "Must", "Results match rate and balance ledgers within tolerance."),
    ("FR-046", "Select on-demand forecast method: target, run-rate, recent average, manual, or zero.", "Must", "Method and assumptions display beside forecast."),
    ("FR-047", "Alert on rapid usage growth, forecast overrun, low balance, imminent expiry, and stale usage data.", "Must", "Each event follows policy, deduplication, and acknowledgement."),
    ("FR-048", "Support vendor API connectors through encrypted credentials/reference when separately approved.", "Could", "Connector failures do not block manual entry and are audited."),
])

add_requirement_group("6.6 Payment Account Management", [
    ("FR-050", "Create card, bank, wallet, virtual card, vendor balance, invoice/manual, reimbursement, and custom payment account types.", "Must", "Required fields adapt to type."),
    ("FR-051", "Store friendly name, provider/issuer, owner, masked identifier/last four, billing currency, expiry month/year, status, limits, notes, and optional icon.", "Must", "No full PAN/CVV/password is accepted or displayed."),
    ("FR-052", "Link a payment account to subscriptions, cost entries, and optional default category/vendor rules.", "Must", "Impact list is shown before status change."),
    ("FR-053", "Alert before account expiry, limit threshold, or inactive-account renewal conflict.", "Must", "Alert includes affected subscriptions and action link."),
    ("FR-054", "Replace or deactivate an account with guided bulk reassignment and effective date.", "Must", "Future charges move; historical charges retain original account."),
    ("FR-055", "Show spend, forecast, subscriptions, next charges, currency exposure, and exceptions per payment account.", "Should", "Totals reconcile with ledger filters."),
])

add_requirement_group("6.7 Targets, Budgets, Forecasts and Alerts", [
    ("FR-060", "Create global, category, tool, subscription, token/on-demand, and payment-account targets.", "Must", "Scope is unambiguous and overlapping targets are allowed with labels."),
    ("FR-061", "Support monthly, quarterly, annual, rolling, custom-date, and one-time target periods.", "Must", "Period boundaries use workspace time zone and fiscal settings."),
    ("FR-062", "Set amount, currency, comparison basis, warning/critical percentages or absolute values, and carry-forward rule.", "Must", "Target utilization is deterministic and displayed."),
    ("FR-063", "Forecast committed and variable cost using actuals, schedules, usage method, trial conversion, and known price changes.", "Must", "Forecast explains components and last-calculated time."),
    ("FR-064", "Create policies for target, renewal, trial, cancellation deadline, payment expiry, stale data, balance, and unusual usage.", "Must", "Policies can be enabled, tested, and scoped."),
    ("FR-065", "Deliver in-app and email alerts with optional future channels, quiet hours, digest, escalation, and recipients.", "Must", "Delivery status and retry history are visible."),
    ("FR-066", "Deduplicate alerts, escalate at higher severity, and allow acknowledge, snooze, dismiss, assign, and resolve.", "Must", "State changes are audited and do not hide unresolved risk."),
    ("FR-067", "Provide notification test and preview using sample data.", "Should", "No real alert is sent unless user confirms test delivery."),
    ("FR-068", "Recommend a target from historical actuals plus user-selected tolerance without changing it automatically.", "Could", "Recommendation shows method, period, and confidence limitations."),
])

add_requirement_group("6.8 Dashboard, Search, Calendar and Reporting", [
    ("FR-070", "Show dashboard KPIs for actual, committed, forecast, target variance, active tools, upcoming renewals, trials, alerts, and payment risks.", "Must", "Cards use the same filter/period context and drill down."),
    ("FR-071", "Show cost trend, category mix, subscription-type mix, top tools, payment-source mix, token trend, and savings opportunities.", "Must", "Charts include accessible table/text alternatives."),
    ("FR-072", "Provide global search and list filters for status, type, category, vendor, owner, payment account, date, amount, tag, alert, and data freshness.", "Must", "Filters combine, persist in URL/state, and can be cleared."),
    ("FR-073", "Provide list, compact card, calendar, and timeline views for subscriptions and renewals.", "Should", "View choice persists and remains usable on mobile."),
    ("FR-074", "Drill from totals to underlying subscriptions, entries, usage, and calculation explanation.", "Must", "Every displayed total is explainable."),
    ("FR-075", "Export filtered data and standard reports to CSV/XLSX and printable/PDF-ready view.", "Should", "Export matches filter context and masks protected fields."),
    ("FR-076", "Schedule periodic email digests and reports.", "Could", "Schedule can be paused and delivery logged."),
    ("FR-077", "Allow saved views with name, filters, columns, sort, density, visibility, and default flag.", "Should", "Users can restore and delete their saved views."),
])

add_requirement_group("6.9 Administration, Audit, Import and Operations", [
    ("FR-080", "Manage users, roles, workspace settings, reference data, feature flags, templates, and supported currencies.", "Must", "Admin actions are permission-checked and audited."),
    ("FR-081", "Maintain immutable audit events for sign-in, CRUD, status, plan, cost, target, account, alert, import, export, and permission changes.", "Must", "Audit includes actor, time, action, entity, before/after summary, and context."),
    ("FR-082", "Provide audit search/export with date, actor, entity, action, outcome, and correlation ID filters.", "Should", "Results respect retention and access rules."),
    ("FR-083", "Provide onboarding wizard and sample-data option with safe removal.", "Should", "A new user can reach a useful dashboard quickly."),
    ("FR-084", "Provide import templates, mapping profiles, validation preview, result summary, and rollback for failed batch.", "Should", "Import results are reproducible and errors actionable."),
    ("FR-085", "Expose background job status for schedule generation, forecast, alerts, email, import, export, and connectors.", "Must", "Failures surface to administrators with retry controls."),
    ("FR-086", "Support logical backup/restore procedures and environment health checks.", "Must", "Recovery is tested and documented."),
    ("FR-087", "Provide configurable data retention and purge workflow with legal/financial hold exceptions.", "Should", "Purge preview identifies affected records and dependencies."),
])

# Document control and contents
add_heading("Document Control", 1)
add_table(
    ["Version", "Date", "Owner", "Change summary", "Status"],
    [
        ("1.0", "29 Jul 2026", "Product Owner", "Initial complete BRD baseline", "For review"),
    ],
    [850, 1250, 1600, 4250, 1410],
    font_size=8.5,
)
add_heading("Approval Record", 2)
add_table(
    ["Role", "Name", "Decision", "Date", "Comments"],
    [
        ("Business Sponsor", "TBD", "Pending", "TBD", ""),
        ("Product Owner", "TBD", "Pending", "TBD", ""),
        ("Technical Lead", "TBD", "Pending", "TBD", ""),
        ("QA Lead", "TBD", "Pending", "TBD", ""),
    ],
    [1800, 1750, 1300, 1200, 3310],
    font_size=8.5,
)
add_heading("Contents", 1)
contents = [
    "1. Executive Summary",
    "2. Business Context, Objectives and Success Measures",
    "3. Scope, Stakeholders, Assumptions and Constraints",
    "4. Business Concepts, Rules and Cost Calculations",
    "5. Future-State Processes",
    "6. Functional Requirements",
    "7. Detailed Use Cases",
    "8. User Experience and Responsive UI Requirements",
    "9. Data Requirements and Logical Data Model",
    "10. Solution Architecture and Technical Requirements",
    "11. Security, Privacy, Audit and Compliance",
    "12. Non-Functional Requirements",
    "13. Reporting, Analytics and Notification Catalogue",
    "14. Testing Strategy and Acceptance Test Scenarios",
    "15. Requirements Traceability",
    "16. Delivery Phasing, Migration and Rollout",
    "17. Risks, Dependencies, Decisions and Sign-off",
    "Appendices: Glossary, Status Models and Sample Calculations",
]
add_compact_contents(contents)
add_callout(
    "DOCUMENT USE",
    "Business stakeholders approve scope and rules; UX teams design against screen requirements; engineering implements "
    "the numbered requirements; QA derives tests from use cases and the traceability matrix.",
    fill=LIGHT_GRAY,
)

page_break()
